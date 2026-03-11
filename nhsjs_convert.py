#!/usr/bin/env python3
"""
nhsjs_convert.py  v4
====================
Convert a LaTeX manuscript (single .tex file, folder, or .zip) to an
NHSJS online-publication Word (.docx) file.

Usage
-----
    python nhsjs_convert.py input.tex [output.docx]
    python nhsjs_convert.py input.zip [output.docx]
    python nhsjs_convert.py project_folder/ [output.docx]

What it produces (NHSJS File 2 requirements)
--------------------------------------------
  * All LaTeX markup stripped; section commands become plain headings,
    environments become native Word constructs.
  * Inline ($...$) and display math kept verbatim as plain text.
  * Figures are EMBEDDED from the project folder (png/pdf/jpg/eps).
  * Figures and tables are AUTO-NUMBERED (Figure 1, Figure 2, ...).
  * Cross-references (\\ref, \\autoref, \\cref, etc.) resolve to the
    correct figure/table/section number instead of [?].
  * Every citation replaced with NHSJS online-citation format:
        ((Author. Title. Journal. Vol. X, pg. Y-Z, Year, DOI.))
    Multiple citations each get their own (( )), separated by a
    superscript comma.  Handles BOTH:
      - \\nhsjsref{1,7,8}  (numbered references in \\begin{enumerate})
      - \\cite{key}        (key-based \\begin{thebibliography})
  * tabular environments converted to Word tables.
  * figure environments: image embedded + numbered caption below.
  * Author block absent from output (NHSJS anonymous review requirement).

Requirements
------------
    pip install python-docx Pillow
"""

import re
import sys
import os
import zipfile
import tempfile
import shutil
import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    print("WARNING: Pillow not installed. Figures will be embedded at default size.",
          file=sys.stderr)


# ============================================================================
# 0.  PROJECT LOADER  -  zip / folder / single .tex
# ============================================================================

def load_project(input_path: str) -> tuple:
    """
    Returns (tex_source, project_root).
    project_root is the directory that contains the figures folder, etc.
    """
    p = Path(input_path)

    if p.is_file() and p.suffix.lower() == '.zip':
        tmp = Path(tempfile.mkdtemp(prefix='nhsjs_'))
        with zipfile.ZipFile(p) as zf:
            zf.extractall(tmp)
        return _find_main_tex(tmp)

    if p.is_dir():
        return _find_main_tex(p)

    if p.is_file() and p.suffix.lower() == '.tex':
        return p.read_text(encoding='utf-8'), p.parent

    raise FileNotFoundError(f"Cannot open input: {input_path!r}")


def _find_main_tex(root: Path) -> tuple:
    """
    Heuristically pick the main .tex file from a directory tree.
    """
    tex_files = sorted(root.rglob('*.tex'))
    if not tex_files:
        raise FileNotFoundError(f"No .tex files found under {root}")

    candidates = [f for f in tex_files
                  if r'\begin{document}' in f.read_text(encoding='utf-8', errors='replace')]
    if not candidates:
        candidates = tex_files

    candidates.sort(key=lambda f: (len(f.parts), -f.stat().st_size))
    chosen = candidates[0]
    print(f"  Main .tex file: {chosen.relative_to(root)}")
    return chosen.read_text(encoding='utf-8'), chosen.parent


# ============================================================================
# 1.  BIBLIOGRAPHY  -  supports BOTH formats
# ============================================================================

def parse_bibliography(tex: str) -> tuple:
    r"""
    Returns:
        num_bib  : {int -> plain_text_citation}   for \nhsjsref / numbered
        key_bib  : {str -> plain_text_citation}   for \cite{key}
    """
    num_bib = {}
    key_bib = {}

    # --- A: enumerate-based numbered references (NHSJS style) ----------------
    enum_m = re.search(
        r'\\section\*?\{References?\}.*?\\begin\{enumerate\}(.*?)\\end\{enumerate\}',
        tex, re.DOTALL | re.IGNORECASE
    )
    if not enum_m:
        enum_m = re.search(
            r'\\begin\{enumerate\}(.*?)\\end\{enumerate\}(?!.*\\begin\{enumerate\})',
            tex, re.DOTALL
        )

    if enum_m:
        raw_enum = enum_m.group(1)
        items = re.split(r'\\item\b', raw_enum)
        for idx, item in enumerate(items):
            item = item.strip()
            if not item:
                continue
            plain = _clean_bib_entry(item)
            if plain:
                num_bib[idx] = plain

    # --- B: thebibliography-based key references -----------------------------
    thebib_m = re.search(
        r'\\begin\{thebibliography\}.*?\\end\{thebibliography\}',
        tex, re.DOTALL
    )
    if thebib_m:
        raw = thebib_m.group(0)
        parts = re.split(r'\\bibitem(?:\[[^\]]*\])?\{([^}]+)\}', raw)
        for i in range(1, len(parts), 2):
            key = parts[i].strip()
            body = parts[i + 1]
            body = re.sub(r'\\end\{thebibliography\}.*', '', body, flags=re.DOTALL)
            plain = _clean_bib_entry(body)
            if plain:
                key_bib[key] = plain

    return num_bib, key_bib


def _clean_bib_entry(s: str) -> str:
    r"""Strip LaTeX markup from a bibliography entry to produce plain text."""
    s = re.sub(r'%[^\n]*', '', s)
    s = re.sub(r'\\url\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\href\{([^}]*)\}\{[^}]*\}', r'\1', s)
    s = re.sub(r'\\texorpdfstring\{([^}]*)\}\{[^}]*\}', r'\1', s)
    for cmd in ('textit', 'emph', 'textbf', 'texttt', 'textrm', 'textnormal',
                'text', 'mathrm'):
        s = re.sub(rf'\\{cmd}\{{([^}}]*)\}}', r'\1', s)
    s = re.sub(r"\\['`\"^~=.]?\{([^}])\}", r'\1', s)
    s = re.sub(r"\\['`\"^~=.]([a-zA-Z])", r'\1', s)
    s = re.sub(r'\\H\{([^}])\}', r'\1', s)
    s = s.replace('---', '\u2014').replace('--', '\u2013')
    s = s.replace('\\ldots', '\u2026').replace('\\dots', '\u2026')
    s = s.replace('\\&', '&').replace('\\%', '%').replace('\\_', '_')
    s = s.replace('\\$', '$').replace('\\#', '#')
    s = re.sub(r'\\[a-zA-Z]+\*?\{[^}]*\}', '', s)
    s = re.sub(r'\\[a-zA-Z]+\*?\b\s*', '', s)
    s = re.sub(r'[{}]', '', s)
    s = re.sub(r'\s+', ' ', s).strip()
    if s and not s.endswith('.'):
        s += '.'
    return s


# ============================================================================
# 2.  MATH PROTECTION
# ============================================================================

_MATH_STORE = []


def protect_math(tex: str) -> str:
    global _MATH_STORE
    _MATH_STORE = []

    def _store(m):
        idx = len(_MATH_STORE)
        _MATH_STORE.append(m.group(0))
        return f'XMATHX{idx:04d}X'

    tex = re.sub(
        r'\\begin\{(equation|align|gather|multline|eqnarray|split)\*?\}'
        r'.*?'
        r'\\end\{\1\*?\}',
        _store, tex, flags=re.DOTALL
    )
    tex = re.sub(r'\\\[.*?\\\]', _store, tex, flags=re.DOTALL)
    tex = re.sub(r'\$\$.*?\$\$', _store, tex, flags=re.DOTALL)
    tex = re.sub(r'\\\(.*?\\\)', _store, tex, flags=re.DOTALL)
    tex = re.sub(r'(?<!\$)\$(?!\$).*?(?<!\$)\$(?!\$)', _store, tex, flags=re.DOTALL)
    return tex


def restore_math_str(s: str) -> str:
    for i, m in enumerate(_MATH_STORE):
        s = s.replace(f'XMATHX{i:04d}X', m)
    return s


# ============================================================================
# 3.  CITATION EXPANSION -> NHSJS online-publication format
# ============================================================================

def expand_nhsjsref(numbers_str: str, num_bib: dict) -> list:
    nums = [n.strip() for n in numbers_str.split(',')]
    result = []
    for idx, n in enumerate(nums):
        try:
            key = int(n)
        except ValueError:
            key = -1
        citation = num_bib.get(key, f'CITATION NOT FOUND: ref {n}')
        if idx > 0:
            result.append(_run(',', superscript=True))
        result.append(_run(f' (({citation}))'))
    return result


def expand_cite(keys_str: str, key_bib: dict) -> list:
    keys = [k.strip() for k in keys_str.split(',')]
    result = []
    for idx, key in enumerate(keys):
        citation = key_bib.get(key, f'CITATION NOT FOUND: {key}')
        if idx > 0:
            result.append(_run(',', superscript=True))
        result.append(_run(f' (({citation}))'))
    return result


# ============================================================================
# 4.  BLOCK DATA CLASSES
# ============================================================================

class TitleBlock:
    def __init__(self, text):
        self.text = text

class HeadingBlock:
    def __init__(self, level, text, label=''):
        self.level = level
        self.text = text
        self.label = label

class ParaBlock:
    def __init__(self, content):
        self.content = content

class TableBlock:
    def __init__(self, caption, rows, label=''):
        self.caption = caption
        self.rows = rows
        self.label = label
        self.number = 0  # assigned after parsing

class FigureBlock:
    def __init__(self, caption, filename='', label=''):
        self.caption = caption
        self.filename = filename
        self.label = label
        self.number = 0  # assigned after parsing

class ListBlock:
    def __init__(self, items, ordered=False):
        self.items = items
        self.ordered = ordered


# ============================================================================
# 5.  BODY EXTRACTION AND BLOCK PARSING
# ============================================================================

def extract_preamble_title(tex: str) -> str:
    """Extract \\title{...} from the preamble (before \\begin{document})."""
    preamble_m = re.search(r'^(.*?)\\begin\{document\}', tex, re.DOTALL)
    preamble = preamble_m.group(1) if preamble_m else ''
    title_m = re.search(r'\\title\{((?:[^{}]|\{(?:[^{}]|\{[^{}]*\})*\})*)\}', preamble)
    if title_m:
        raw = title_m.group(1)
        raw = re.sub(r'\\\\(\[[^\]]*\])?', ' ', raw).strip()
        return raw
    return ''


def get_body(tex: str) -> str:
    m = re.search(r'\\begin\{document\}(.*?)\\end\{document\}', tex, re.DOTALL)
    body = m.group(1) if m else tex

    body = re.sub(
        r'\\begin\{thebibliography\}.*?\\end\{thebibliography\}',
        '', body, flags=re.DOTALL
    )
    body = re.sub(
        r'\\section\*?\{References?\}.*?\\begin\{enumerate\}.*?\\end\{enumerate\}',
        '', body, flags=re.DOTALL | re.IGNORECASE
    )
    return body


def parse_body(body: str, preamble_title: str = '') -> list:
    body = re.sub(r'%[^\n]*', '', body)
    body = re.sub(r'\\appendix\b', r'\n\\section{Appendix}\n', body)

    _layout = (r'FloatBarrier|clearpage|newpage|vfill|hfill|'
               r'raggedright|raggedleft|centering|'
               r'setstretch\{[^}]*\}|medskip|bigskip|smallskip|'
               r'setlength\{[^}]*\}\{[^}]*\}|sloppy|noindent')
    body = re.sub(rf'\\(?:{_layout})\b', '', body)
    body = re.sub(r'\\(?:vspace|hspace)\*?\{[^}]*\}', '', body)

    events = []

    # --- Title: check body first, fall back to preamble ---
    title_from_cmd = re.search(r'\\title\{((?:[^{}]|\{[^{}]*\})*)\}', body)
    maketitle_m = re.search(r'\\maketitle\b', body)

    title_text = ''
    if title_from_cmd:
        title_text = title_from_cmd.group(1)
        title_text = re.sub(r'\\\\(\[[^\]]*\])?', ' ', title_text).strip()
        body = body[:title_from_cmd.start()] + ' ' * (title_from_cmd.end() - title_from_cmd.start()) + body[title_from_cmd.end():]
    elif preamble_title:
        title_text = preamble_title

    if title_text:
        if maketitle_m:
            events.append((maketitle_m.start(), maketitle_m.end(),
                           TitleBlock(title_text)))
        else:
            # No \maketitle, insert title at the very beginning
            events.append((0, 0, TitleBlock(title_text)))

    # Flatten center blocks and minipages
    body = re.sub(r'\\begin\{center\}(.*?)\\end\{center\}',
                  r'\1', body, flags=re.DOTALL)
    body = re.sub(r'\\begin\{minipage\}(?:\[[^\]]*\])?\{[^}]*\}(.*?)\\end\{minipage\}',
                  r'\1', body, flags=re.DOTALL)

    # --- Section headings ---
    for m in re.finditer(
            r'\\(section|subsection|subsubsection)\*?\{((?:[^{}]|\{[^{}]*\})*)\}',
            body):
        level = {'section': 1, 'subsection': 2, 'subsubsection': 3}[m.group(1)]
        text = _heading_text(m.group(2))
        # Look for a \label right after the heading
        label = ''
        after = body[m.end():m.end()+100]
        lm = re.match(r'\s*\\label\{([^}]+)\}', after)
        if lm:
            label = lm.group(1).strip()
        events.append((m.start(), m.end(), HeadingBlock(level, text, label)))

    for m in re.finditer(r'\\paragraph\{((?:[^{}]|\{[^{}]*\})*)\}', body):
        events.append((m.start(), m.end(),
                       HeadingBlock(3, _heading_text(m.group(1)))))

    # --- Tables ---
    for m in re.finditer(r'\\begin\{table\*?\}.*?\\end\{table\*?\}',
                         body, re.DOTALL):
        events.append((m.start(), m.end(), _parse_table(m.group(0))))

    # --- Figures ---
    for m in re.finditer(r'\\begin\{figure\*?\}.*?\\end\{figure\*?\}',
                         body, re.DOTALL):
        events.append((m.start(), m.end(), _parse_figure(m.group(0))))

    # --- Standalone \includegraphics outside figure env ---
    for m in re.finditer(
            r'(?<!\\begin\{figure)\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}',
            body):
        # Check it's not inside a figure environment already captured
        inside = False
        for start, end, blk in events:
            if isinstance(blk, FigureBlock) and start <= m.start() < end:
                inside = True
                break
        if not inside:
            fb = FigureBlock(caption='', filename=m.group(1).strip(), label='')
            events.append((m.start(), m.end(), fb))

    # --- Lists ---
    for m in re.finditer(r'\\begin\{itemize\}(.*?)\\end\{itemize\}',
                         body, re.DOTALL):
        items = _split_items(m.group(1))
        events.append((m.start(), m.end(), ListBlock(items, ordered=False)))

    for m in re.finditer(r'\\begin\{enumerate\}(.*?)\\end\{enumerate\}',
                         body, re.DOTALL):
        items = _split_items(m.group(1))
        events.append((m.start(), m.end(), ListBlock(items, ordered=True)))

    events.sort(key=lambda x: x[0])
    events = _remove_overlaps(events)

    blocks = []
    pos = 0
    for start, end, block in events:
        _add_text_paras(body[pos:start], blocks)
        blocks.append(block)
        pos = end
    _add_text_paras(body[pos:], blocks)
    return blocks


def _heading_text(raw: str) -> str:
    raw = re.sub(r'\\texorpdfstring\{((?:[^{}]|\{[^{}]*\})*)\}\{[^}]*\}',
                 r'\1', raw)
    for cmd in ('textit', 'emph', 'textbf', 'texttt', 'textrm', 'textnormal',
                'text'):
        raw = re.sub(rf'\\{cmd}\{{((?:[^{{}}]|\{{[^{{}}]*\}})*)\}}', r'\1', raw)
    raw = raw.replace('\\ ', ' ').replace('\\,', ' ').replace('\\;', ' ')
    raw = raw.replace('\\&', '&').replace('\\%', '%').replace('\\_', '_')
    raw = re.sub(r'\\label\{[^}]*\}', '', raw)
    raw = re.sub(r'\\[a-zA-Z]+\*?\b\s*', '', raw)
    raw = re.sub(r'[{}]', '', raw)
    raw = restore_math_str(raw)
    return raw.strip()


def _remove_overlaps(events):
    result, last_end = [], -1
    for start, end, block in events:
        if start >= last_end:
            result.append((start, end, block))
            last_end = end
    return result


def _split_items(text: str) -> list:
    raw_items = re.split(r'\\item\b', text)
    result = []
    for item in raw_items:
        item = item.strip()
        item = re.sub(r'^\[[^\]]*\]\s*', '', item)
        if item:
            result.append(item)
    return result


def _add_text_paras(text: str, blocks: list):
    for chunk in re.split(r'\n\s*\n', text):
        chunk = chunk.strip()
        if not chunk:
            continue
        if re.match(r'^\\(?:label|begin|end|graphicspath|newcommand|'
                    r'renewcommand|setlength|usepackage|documentclass|'
                    r'geometry|maketitle|date|title|author)\s*[\[{]?', chunk):
            continue
        stripped = re.sub(r'\\[a-zA-Z]+\*?|\{|\}|\[.*?\]', '', chunk).strip()
        if not stripped:
            continue
        blocks.append(ParaBlock(chunk))


def _parse_table(env: str) -> TableBlock:
    caption = _extract_env_caption(env)
    label = ''
    lm = re.search(r'\\label\{([^}]+)\}', env)
    if lm:
        label = lm.group(1).strip()

    tab_m = re.search(r'\\begin\{tabular[x*]?\}.*?\\end\{tabular[x*]?\}',
                      env, re.DOTALL)
    rows = []
    if tab_m:
        tab = tab_m.group(0)
        tab = re.sub(r'\\begin\{tabular[x*]?\}\{[^}]*\}', '', tab)
        tab = re.sub(r'\\end\{tabular[x*]?\}', '', tab)
        tab = re.sub(r'\\(?:toprule|midrule|bottomrule|hline|cline\{[^}]*\})\b',
                     '', tab)
        for raw_row in re.split(r'\\\\(?:\[[^\]]*\])?', tab):
            raw_row = raw_row.strip()
            if not raw_row:
                continue
            cells = []
            for c in raw_row.split('&'):
                c = restore_math_str(c.strip())
                c = _strip_cell(c)
                cells.append(c.strip())
            if any(c for c in cells):
                rows.append(cells)
    return TableBlock(caption, rows, label)


def _strip_cell(s: str) -> str:
    s = s.replace('\\&', '&').replace('\\_', '_').replace('\\%', '%')
    s = s.replace('\\$', '$').replace('\\pm', '\u00b1')
    for cmd in ('textbf', 'textit', 'emph', 'text', 'textrm', 'texttt',
                'multicolumn', 'multirow'):
        s = re.sub(rf'\\{cmd}\{{[^}}]*\}}\{{([^}}]*)\}}', r'\1', s)
        s = re.sub(rf'\\{cmd}\{{([^}}]*)\}}', r'\1', s)
    s = re.sub(r'\\[a-zA-Z]+\*?\{[^}]*\}', '', s)
    s = re.sub(r'\\[a-zA-Z]+\b', '', s)
    s = re.sub(r'[{}]', '', s)
    return s.strip()


def _parse_figure(env: str) -> FigureBlock:
    caption = _extract_env_caption(env)
    label = ''
    lm = re.search(r'\\label\{([^}]+)\}', env)
    if lm:
        label = lm.group(1).strip()
    filename = ''
    for pattern in (
        r'\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}',
        r'\\includefig(?:\[[^\]]*\])?\{([^}]+)\}',
    ):
        m = re.search(pattern, env)
        if m:
            filename = m.group(1).strip()
            break
    return FigureBlock(caption, filename, label)


def _extract_env_caption(env: str) -> str:
    m = re.search(r'\\caption(?:\[[^\]]*\])?\{((?:[^{}]|\{(?:[^{}]|\{[^{}]*\})*\})*)\}', env)
    if not m:
        return ''
    raw = m.group(1)
    raw = re.sub(r'\\texorpdfstring\{((?:[^{}]|\{[^{}]*\})*)\}\{[^}]*\}',
                 r'\1', raw)
    return raw.strip()


# ============================================================================
# 5b. LABEL MAP BUILDER  -  assigns numbers and resolves cross-references
# ============================================================================

def build_label_map(blocks: list) -> dict:
    """
    Walk through blocks in order, assign sequential numbers to figures
    and tables, and build a mapping: {label_string -> display_string}.

    Also assigns .number to each FigureBlock and TableBlock.

    Handles section numbering too for section labels.
    """
    label_map = {}
    fig_counter = 0
    tbl_counter = 0
    section_counters = [0, 0, 0]  # section, subsection, subsubsection

    for block in blocks:
        if isinstance(block, FigureBlock):
            fig_counter += 1
            block.number = fig_counter
            if block.label:
                label_map[block.label] = str(fig_counter)
                # For \autoref-style: store the "Figure X" form too
                label_map[f'__autoref__{block.label}'] = f'Figure {fig_counter}'

        elif isinstance(block, TableBlock):
            tbl_counter += 1
            block.number = tbl_counter
            if block.label:
                label_map[block.label] = str(tbl_counter)
                label_map[f'__autoref__{block.label}'] = f'Table {tbl_counter}'

        elif isinstance(block, HeadingBlock):
            if block.level == 1:
                section_counters[0] += 1
                section_counters[1] = 0
                section_counters[2] = 0
                sec_num = str(section_counters[0])
            elif block.level == 2:
                section_counters[1] += 1
                section_counters[2] = 0
                sec_num = f'{section_counters[0]}.{section_counters[1]}'
            else:
                section_counters[2] += 1
                sec_num = f'{section_counters[0]}.{section_counters[1]}.{section_counters[2]}'

            if block.label:
                label_map[block.label] = sec_num
                label_map[f'__autoref__{block.label}'] = f'Section {sec_num}'

    print(f'  Label map: {len(label_map)//2 if label_map else 0} labeled items '
          f'({fig_counter} figures, {tbl_counter} tables)')
    return label_map


# ============================================================================
# 6.  IMAGE RESOLVER
# ============================================================================

IMAGE_EXTENSIONS = ['.png', '.jpg', '.jpeg', '.pdf', '.eps',
                    '.PNG', '.JPG', '.JPEG', '.svg', '.SVG']


def find_image(raw_filename: str, project_root: Path) -> "Path | None":
    if not raw_filename:
        return None

    candidates = [raw_filename]
    for ext in IMAGE_EXTENSIONS:
        candidates.append(raw_filename + ext)

    for c in candidates:
        p = project_root / c
        if p.exists():
            return p

    # Recursive search by basename
    stem = Path(raw_filename).stem
    for ext in IMAGE_EXTENSIONS:
        matches = list(project_root.rglob(stem + ext))
        if matches:
            return matches[0]

    return None


def get_image_dimensions_inches(img_path: Path, max_width_in=6.0) -> tuple:
    if not HAS_PIL:
        return (max_width_in, 4.0)
    try:
        with PILImage.open(img_path) as img:
            w_px, h_px = img.size
            dpi = img.info.get('dpi', (96, 96))
            if isinstance(dpi, (int, float)):
                dpi = (dpi, dpi)
            dpi_x = max(dpi[0], 1)
            dpi_y = max(dpi[1], 1)
            w_in = w_px / dpi_x
            h_in = h_px / dpi_y
            if w_in > max_width_in:
                ratio = max_width_in / w_in
                w_in = max_width_in
                h_in *= ratio
            return (w_in, h_in)
    except Exception:
        return (max_width_in, 4.0)


# ============================================================================
# 7.  INLINE TEXT PROCESSOR -> list of run-spec dicts
# ============================================================================

def process_inline(text: str, num_bib: dict, key_bib: dict,
                   label_map: dict = None) -> list:
    if label_map is None:
        label_map = {}

    # Pre-clean
    text = re.sub(r'\\label\{[^}]*\}', '', text)
    text = re.sub(r'\\(?:FloatBarrier|clearpage|newpage)\b', '', text)
    text = re.sub(r'\\(?:vspace|hspace)\*?\{[^}]*\}', '', text)
    text = re.sub(r'\\\\(\[[^\]]*\])?', '\n', text)
    text = re.sub(r'\\newline\b', '\n', text)
    text = text.replace('\\,', '\u2009').replace('\\;', ' ').replace('\\:', ' ')
    text = text.replace('\\!', '')
    text = text.replace('\\&', '&').replace('\\%', '%')
    text = text.replace('\\textendash', '\u2013').replace('\\textemdash', '\u2014')
    text = text.replace('---', '\u2014').replace('--', '\u2013')
    text = text.replace('\\ldots', '\u2026').replace('\\dots', '\u2026')
    text = text.replace('~', '\u00a0')
    text = re.sub(r'\\keywords\{([^}]*)\}', r'Keywords: \1', text)
    text = re.sub(r'\\maketitle\b', '', text)
    text = re.sub(r'\\date\{[^}]*\}', '', text)
    text = re.sub(r'\\title\{(?:[^{}]|\{[^{}]*\})*\}', '', text)
    text = re.sub(r'\\author\{(?:[^{}]|\{[^{}]*\})*\}', '', text)

    runs = _tokenize(text, num_bib, key_bib, label_map)

    # Merge adjacent runs with identical formatting
    merged = []
    for r in runs:
        if (merged
                and merged[-1]['bold'] == r['bold']
                and merged[-1]['italic'] == r['italic']
                and merged[-1]['superscript'] == r['superscript']):
            merged[-1]['text'] += r['text']
        else:
            merged.append(dict(r))

    if merged and merged[0]['text'].lstrip(' \u00a0'):
        merged[0]['text'] = merged[0]['text'].lstrip(' \u00a0')

    return [r for r in merged if r['text']]


def _run(text, bold=False, italic=False, superscript=False):
    return {'text': text, 'bold': bold, 'italic': italic,
            'superscript': superscript}


def _tokenize(text: str, num_bib: dict, key_bib: dict,
              label_map: dict,
              bold=False, italic=False) -> list:
    runs = []
    i, n = 0, len(text)

    while i < n:

        # Math token
        mm = re.match(r'XMATHX(\d{4})X', text[i:])
        if mm:
            idx = int(mm.group(1))
            math_text = _MATH_STORE[idx] if idx < len(_MATH_STORE) else ''
            runs.append(_run(math_text, bold=bold, italic=italic))
            i += len(mm.group(0))
            continue

        # Literal $
        if text[i] == '$':
            if i + 1 < n and text[i+1] == '$':
                end = text.find('$$', i + 2)
                end = n - 2 if end == -1 else end
                runs.append(_run(text[i:end+2], bold=bold, italic=italic))
                i = end + 2
            else:
                end = text.find('$', i + 1)
                end = n - 1 if end == -1 else end
                runs.append(_run(text[i:end+1], bold=bold, italic=italic))
                i = end + 1
            continue

        # Plain text
        if text[i] != '\\':
            j = i
            while j < n and text[j] not in ('\\', '{', '}', '$'):
                if text[j:j+6] == 'XMATHX':
                    break
                j += 1
            if j > i:
                runs.append(_run(text[i:j], bold=bold, italic=italic))
            i = j
            if i < n and text[i] == '{':
                inner, i = _braced(text, i)
                runs.extend(_tokenize(inner, num_bib, key_bib, label_map, bold, italic))
            elif i < n and text[i] == '}':
                i += 1
            continue

        # Backslash
        if i + 1 < n:
            nc = text[i + 1]
            if nc == '_':
                runs.append(_run('_', bold=bold, italic=italic)); i += 2; continue
            if nc == '-':
                i += 2; continue
            if nc == '{':
                runs.append(_run('{', bold=bold, italic=italic)); i += 2; continue
            if nc == '}':
                runs.append(_run('}', bold=bold, italic=italic)); i += 2; continue
            if nc == ' ':
                runs.append(_run(' ', bold=bold, italic=italic)); i += 2; continue
            if nc in ("'", '`', '"', '^', '~', '=', '.'):
                i += 2
                if i < n and text[i] == '{':
                    inner, i = _braced(text, i)
                    runs.append(_run(inner, bold=bold, italic=italic))
                elif i < n and text[i].isalpha():
                    runs.append(_run(text[i], bold=bold, italic=italic))
                    i += 1
                continue

        m = re.match(r'\\([a-zA-Z]+)\*?', text[i:])
        if not m:
            runs.append(_run(text[i], bold=bold, italic=italic))
            i += 1
            continue

        cmd = m.group(1)
        i += len(m.group(0))

        # --- Formatting commands ---
        if cmd in ('textbf', 'mathbf', 'boldsymbol', 'bm'):
            inner, i = _braced(text, i)
            runs.extend(_tokenize(inner, num_bib, key_bib, label_map, bold=True, italic=italic))

        elif cmd in ('textit', 'emph', 'mathit'):
            inner, i = _braced(text, i)
            runs.extend(_tokenize(inner, num_bib, key_bib, label_map, bold=bold, italic=True))

        elif cmd in ('texttt', 'textsc', 'textrm', 'textnormal', 'text',
                     'mathrm', 'mathcal', 'mathbb', 'textup', 'textsf',
                     'textmd'):
            if i < n and text[i] == '{':
                inner, i = _braced(text, i)
                runs.extend(_tokenize(inner, num_bib, key_bib, label_map, bold=bold, italic=italic))

        elif cmd in ('small', 'normalsize', 'footnotesize', 'large',
                     'Large', 'LARGE', 'huge', 'Huge', 'scriptsize', 'tiny'):
            if i < n and text[i] == '{':
                inner, i = _braced(text, i)
                runs.extend(_tokenize(inner, num_bib, key_bib, label_map, bold=bold, italic=italic))

        elif cmd == 'texorpdfstring':
            latex_form, i = _braced(text, i)
            _, i = _braced(text, i)
            runs.extend(_tokenize(latex_form, num_bib, key_bib, label_map, bold=bold, italic=italic))

        # --- NHSJS numbered citation ---
        elif cmd == 'nhsjsref':
            inner, i = _braced(text, i)
            runs.extend(expand_nhsjsref(inner, num_bib))

        # --- Standard \cite ---
        elif cmd == 'cite':
            if i < n and text[i] == '[':
                _, i = _bracketed(text, i)
            inner, i = _braced(text, i)
            runs.extend(expand_cite(inner, key_bib))

        # --- \citep, \citet, etc. ---
        elif cmd in ('citep', 'citet', 'citealt', 'citealp', 'citeauthor',
                     'citeyear', 'citenum'):
            if i < n and text[i] == '[':
                _, i = _bracketed(text, i)
            if i < n and text[i] == '[':
                _, i = _bracketed(text, i)
            inner, i = _braced(text, i)
            runs.extend(expand_cite(inner, key_bib))

        # --- Cross-references: \ref, \eqref -> just the number ---
        elif cmd in ('ref', 'eqref', 'pageref'):
            inner, i = _braced(text, i)
            resolved = label_map.get(inner.strip(), '??')
            if cmd == 'eqref':
                resolved = f'({resolved})'
            runs.append(_run(resolved, bold=bold, italic=italic))

        # --- \autoref, \cref, \Cref -> "Figure 1", "Table 2", etc. ---
        elif cmd in ('autoref', 'cref', 'Cref', 'nameref'):
            inner, i = _braced(text, i)
            key = inner.strip()
            # Try the autoref form first ("Figure 1"), fall back to number
            resolved = label_map.get(f'__autoref__{key}')
            if resolved is None:
                num = label_map.get(key, '??')
                # Try to guess the type from the label prefix
                prefix = _guess_ref_prefix(key)
                resolved = f'{prefix}{num}' if prefix else num
            runs.append(_run(resolved, bold=bold, italic=italic))

        elif cmd == 'label':
            _, i = _braced(text, i)

        elif cmd == 'url':
            inner, i = _braced(text, i)
            runs.append(_run(inner, bold=bold, italic=italic))

        elif cmd == 'href':
            url, i = _braced(text, i)
            label, i = _braced(text, i)
            runs.append(_run(label if label else url, bold=bold, italic=italic))

        elif cmd == 'footnote':
            inner, i = _braced(text, i)
            inner_runs = _tokenize(inner, num_bib, key_bib, label_map, bold=bold, italic=italic)
            runs.append(_run(' [Note: ', bold=bold, italic=italic))
            runs.extend(inner_runs)
            runs.append(_run(']', bold=bold, italic=italic))

        elif cmd in ('begin', 'end'):
            _, i = _braced(text, i)

        elif cmd in ('newline', 'par'):
            runs.append(_run('\n'))

        elif cmd in ('quad', 'qquad', 'enspace', 'thinspace', 'medspace',
                     'thickspace'):
            runs.append(_run(' '))

        elif cmd in ('hfill', 'hspace', 'vspace'):
            if i < n and text[i] == '{':
                _, i = _braced(text, i)
            runs.append(_run(' '))

        elif cmd in ('noindent', 'centering', 'raggedright', 'raggedleft',
                     'FloatBarrier', 'clearpage', 'newpage',
                     'medskip', 'bigskip', 'smallskip',
                     'linebreak', 'pagebreak', 'allowdisplaybreaks',
                     'maketitle', 'tableofcontents'):
            if i < n and text[i] == '{':
                _, i = _braced(text, i)

        elif cmd == 'xrightarrow':
            if i < n and text[i] == '{':
                _, i = _braced(text, i)
            runs.append(_run('\u2192', bold=bold, italic=italic))

        elif cmd in ('rightarrow', 'to'):
            runs.append(_run('\u2192', bold=bold, italic=italic))
        elif cmd == 'leftarrow':
            runs.append(_run('\u2190', bold=bold, italic=italic))
        elif cmd in ('Rightarrow', 'implies'):
            runs.append(_run('\u21D2', bold=bold, italic=italic))
        elif cmd == 'infty':
            runs.append(_run('\u221e', bold=bold, italic=italic))
        elif cmd == 'times':
            runs.append(_run('\u00d7', bold=bold, italic=italic))
        elif cmd == 'pm':
            runs.append(_run('\u00b1', bold=bold, italic=italic))
        elif cmd == 'leq':
            runs.append(_run('\u2264', bold=bold, italic=italic))
        elif cmd == 'geq':
            runs.append(_run('\u2265', bold=bold, italic=italic))
        elif cmd == 'neq':
            runs.append(_run('\u2260', bold=bold, italic=italic))
        elif cmd == 'approx':
            runs.append(_run('\u2248', bold=bold, italic=italic))
        elif cmd == 'alpha':
            runs.append(_run('\u03b1', bold=bold, italic=italic))
        elif cmd == 'beta':
            runs.append(_run('\u03b2', bold=bold, italic=italic))
        elif cmd == 'gamma':
            runs.append(_run('\u03b3', bold=bold, italic=italic))
        elif cmd == 'delta':
            runs.append(_run('\u03b4', bold=bold, italic=italic))
        elif cmd == 'Delta':
            runs.append(_run('\u0394', bold=bold, italic=italic))
        elif cmd == 'pi':
            runs.append(_run('\u03c0', bold=bold, italic=italic))
        elif cmd == 'sigma':
            runs.append(_run('\u03c3', bold=bold, italic=italic))
        elif cmd == 'mu':
            runs.append(_run('\u03bc', bold=bold, italic=italic))

        elif cmd == 'H':
            if i < n and text[i] == '{':
                inner, i = _braced(text, i)
                runs.append(_run(inner, bold=bold, italic=italic))

        else:
            if i < n and text[i] == '{':
                inner, i = _braced(text, i)
                runs.extend(_tokenize(inner, num_bib, key_bib, label_map, bold=bold, italic=italic))

    return runs


def _guess_ref_prefix(label: str) -> str:
    """Guess a prefix like 'Figure ' from a label like 'fig:myfigure'."""
    label_lower = label.lower()
    if label_lower.startswith('fig:') or label_lower.startswith('fig_'):
        return 'Figure '
    if label_lower.startswith('tab:') or label_lower.startswith('tab_'):
        return 'Table '
    if label_lower.startswith('sec:') or label_lower.startswith('sec_'):
        return 'Section '
    if label_lower.startswith('eq:') or label_lower.startswith('eq_'):
        return 'Eq. '
    return ''


def _braced(text: str, i: int):
    if i >= len(text) or text[i] != '{':
        return '', i
    depth, j = 0, i
    while j < len(text):
        if text[j] == '{':
            depth += 1
        elif text[j] == '}':
            depth -= 1
            if depth == 0:
                return text[i + 1:j], j + 1
        j += 1
    return text[i + 1:], len(text)


def _bracketed(text: str, i: int):
    if i >= len(text) or text[i] != '[':
        return '', i
    j = text.find(']', i)
    if j == -1:
        return '', i
    return text[i + 1:j], j + 1


# ============================================================================
# 8.  DOCX WRITER
# ============================================================================

def _add_runs_to_para(para, runs: list):
    for r in runs:
        if not r['text']:
            continue
        run = para.add_run(r['text'])
        run.bold = r['bold']
        run.italic = r['italic']
        if r['superscript']:
            run.font.superscript = True


def write_docx(blocks: list, num_bib: dict, key_bib: dict,
               label_map: dict, project_root: Path, output_path: str):
    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width = int(8.5 * 914400)
    sec.page_height = int(11 * 914400)
    for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, attr, Inches(1))

    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    # Force all heading styles to 12pt Times New Roman, bold
    for lvl in range(1, 4):
        style_name = f'Heading {lvl}'
        try:
            h_style = doc.styles[style_name]
            h_style.font.name = 'Times New Roman'
            h_style.font.size = Pt(12)
            h_style.font.bold = True
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(4)
        except (KeyError, ValueError):
            pass

    for block in blocks:

        # -- Title --
        if isinstance(block, TitleBlock):
            runs = process_inline(block.text, num_bib, key_bib, label_map)
            if not runs:
                continue
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in runs:
                if not r['text']:
                    continue
                run = para.add_run(r['text'])
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Times New Roman'

        # -- Heading --
        elif isinstance(block, HeadingBlock):
            text = block.text.strip()
            if not text:
                continue
            doc.add_heading(text, level=min(block.level, 3))

        # -- Paragraph --
        elif isinstance(block, ParaBlock):
            content = block.content.strip()
            if not content:
                continue
            runs = process_inline(content, num_bib, key_bib, label_map)
            if not runs:
                continue
            para = doc.add_paragraph()
            _add_runs_to_para(para, runs)

        # -- Table --
        elif isinstance(block, TableBlock):
            if block.caption:
                cap_text = restore_math_str(block.caption)
                cap_runs = process_inline(cap_text, num_bib, key_bib, label_map)
                cap = doc.add_paragraph()
                num_str = f'Table {block.number}: ' if block.number else 'Table: '
                bold_run = cap.add_run(num_str)
                bold_run.bold = True
                bold_run.font.name = 'Times New Roman'
                bold_run.italic = True
                _add_runs_to_para(cap, cap_runs)

            if not block.rows:
                doc.add_paragraph('[TABLE: could not parse rows]')
                continue

            ncols = max(len(row) for row in block.rows)
            tbl = doc.add_table(rows=len(block.rows), cols=ncols)
            tbl.style = 'Table Grid'
            for ri, row in enumerate(block.rows):
                for ci in range(ncols):
                    cell_text = row[ci] if ci < len(row) else ''
                    cell = tbl.cell(ri, ci)
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    if ri == 0:
                        run.bold = True
            doc.add_paragraph()

        # -- Figure --
        elif isinstance(block, FigureBlock):
            img_path = find_image(block.filename, project_root)

            if img_path and img_path.suffix.lower() not in ('.pdf', '.eps'):
                try:
                    w_in, h_in = get_image_dimensions_inches(img_path)
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(str(img_path), width=Inches(w_in))
                    print(f"  Embedded figure {block.number}: {img_path.name}")
                except Exception as e:
                    print(f"  WARNING: Could not embed {img_path}: {e}",
                          file=sys.stderr)
                    _add_figure_placeholder(doc, block.filename, block.number)
            elif img_path and img_path.suffix.lower() in ('.pdf', '.eps'):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f'[FIGURE {block.number}: {img_path.name} \u2014 PDF/EPS, see supplementary]')
                r.italic = True
                r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            else:
                _add_figure_placeholder(doc, block.filename, block.number)

            # Caption with number
            if block.caption:
                cap_raw = restore_math_str(block.caption)
                cap_runs = process_inline(cap_raw, num_bib, key_bib, label_map)
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                num_str = f'Figure {block.number}: ' if block.number else 'Figure: '
                bold_run = cp.add_run(num_str)
                bold_run.bold = True
                bold_run.italic = True
                bold_run.font.name = 'Times New Roman'
                _add_runs_to_para(cp, cap_runs)
            elif block.number:
                # No caption text, but still show "Figure N" below image
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                bold_run = cp.add_run(f'Figure {block.number}')
                bold_run.bold = True
                bold_run.italic = True
                bold_run.font.name = 'Times New Roman'

            doc.add_paragraph()

        # -- List --
        elif isinstance(block, ListBlock):
            list_style = 'List Number' if block.ordered else 'List Bullet'
            for item in block.items:
                if not item.strip():
                    continue
                runs = process_inline(item, num_bib, key_bib, label_map)
                if not runs:
                    continue
                para = doc.add_paragraph(style=list_style)
                _add_runs_to_para(para, runs)

    # -- Append References section at end --
    if num_bib or key_bib:
        doc.add_heading('References', level=1)

        if num_bib:
            # Numbered references (NHSJS enumerate style)
            sorted_keys = sorted(num_bib.keys())
            for idx, key in enumerate(sorted_keys, start=1):
                para = doc.add_paragraph(style='List Number')
                run = para.add_run(num_bib[key])
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        elif key_bib:
            # Key-based references (thebibliography style)
            for idx, (key, text) in enumerate(key_bib.items(), start=1):
                para = doc.add_paragraph(style='List Number')
                run = para.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    doc.save(output_path)
    print(f'\nSaved: {output_path}')


def _add_figure_placeholder(doc, filename: str, number: int = 0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    num_part = f' {number}' if number else ''
    if filename:
        label = f'[FIGURE{num_part} NOT FOUND: {filename}]'
    else:
        label = f'[FIGURE{num_part} PLACEHOLDER]'
    r = p.add_run(label)
    r.italic = True
    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


# ============================================================================
# 9.  MAIN
# ============================================================================

def main():
    parser = argparse.ArgumentParser(
        description='Convert a LaTeX project to NHSJS online-publication Word format.',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    parser.add_argument('input',
        help='Input: a .tex file, a .zip of the Overleaf project, or a folder.')
    parser.add_argument('output', nargs='?',
        help='Output .docx path (default: <input stem>_nhsjs_online.docx).')
    args = parser.parse_args()

    in_path = Path(args.input)
    out_stem = in_path.stem if in_path.is_file() else in_path.name
    out_path = Path(args.output) if args.output \
               else Path(out_stem + '_nhsjs_online.docx')

    print(f'Loading project from: {in_path}')
    tex, project_root = load_project(str(in_path))
    print(f'  Project root: {project_root}')

    # 1. Parse bibliography BEFORE math protection
    num_bib, key_bib = parse_bibliography(tex)
    print(f'  Numbered refs: {len(num_bib)},  Key-based refs: {len(key_bib)}')
    if not num_bib and not key_bib:
        print('  WARNING: No references found.', file=sys.stderr)

    # 2. Protect math
    tex = protect_math(tex)

    # 3. Extract preamble title (before \begin{document})
    preamble_title = extract_preamble_title(tex)
    if preamble_title:
        print(f'  Title from preamble: {preamble_title[:60]}...'
              if len(preamble_title) > 60 else f'  Title from preamble: {preamble_title}')

    # 4. Extract body
    body = get_body(tex)

    # 5. Parse blocks
    blocks = parse_body(body, preamble_title)
    counts = {}
    for b in blocks:
        k = type(b).__name__
        counts[k] = counts.get(k, 0) + 1
    print(f'  Blocks: {counts}')

    # 5. Build label map (assigns figure/table numbers, resolves cross-refs)
    label_map = build_label_map(blocks)

    # 6. Write docx
    write_docx(blocks, num_bib, key_bib, label_map, project_root, str(out_path))


if __name__ == '__main__':
    main()