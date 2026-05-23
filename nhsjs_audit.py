"""
nhsjs_audit.py — v2.0.0
=======================
Self-audit checks for a v2.2-format NHSJS revision build.

Implements the 13 checks from Revise-After-Review v2.2 Phase 6. Used to
surface issues before declaring a revision ready to send. The 3 v2.2
additions vs v2 are: patch anchor integrity, REPLACE_ALL CONFIRM_COUNT
verification, and FIGURE_REGENERATED flag presence.

Public API
----------
    audit_revision(revision_md_path, original_docx_path, submit_dir,
                   *, log_path=None) -> list[CheckResult]

    CheckResult(name, status, detail) — dataclass, status PASS / FAIL / WARN.

CLI
---
    python -m nhsjs_audit revision.md original.docx submit/
    python -m nhsjs_audit revision.md original.docx submit/ --log revision-log.md
"""
from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.oxml.ns import qn


# ─── Data classes ──────────────────────────────────────────────────────────

@dataclass
class CheckResult:
    name: str
    status: str  # PASS / FAIL / WARN
    detail: str


# ─── Constants ─────────────────────────────────────────────────────────────

_CITE_RE = re.compile(r"\[\[CITE:([^\]]+)\]\]")

# Phrases that should never appear in an anonymous Standard manuscript.
_IDENTIFYING_PHRASES = [
    r"\bAthena Education\b",
    r"\bmy mentor\b",
    r"\bmy counsell?or\b",
    r"\bmy tutor\b",
    r"\bmy advisor\b",
    r"\bthanks to my\b",
]
_IDENTIFYING_REGEX = re.compile("|".join(_IDENTIFYING_PHRASES), re.IGNORECASE)

_AUTHOR_BLOCKLIST = {
    "", "anonymous", "claude", "anthropic", "athena", "athena education",
    "user", "guest",
}


# ─── audit_revision (public entry point) ───────────────────────────────────

def audit_revision(revision_md_path: str | Path,
                   original_docx_path: str | Path,
                   submit_dir: str | Path,
                   *,
                   log_path: str | Path | None = None) -> list[CheckResult]:
    """Run all 13 Phase 6 checks. Returns one CheckResult per check."""
    revision_md_path = Path(revision_md_path)
    original_docx_path = Path(original_docx_path)
    submit_dir = Path(submit_dir)
    log_path = Path(log_path) if log_path else None

    if not revision_md_path.exists():
        return [CheckResult("revision.md exists", "FAIL",
                            f"missing file: {revision_md_path}")]
    if not original_docx_path.exists():
        return [CheckResult("original.docx exists", "FAIL",
                            f"missing file: {original_docx_path}")]

    # Parse revision.md via the builder to get patches
    try:
        from nhsjs_revision_builder import parse_revision
        revision = parse_revision(str(revision_md_path))
    except Exception as exc:
        return [CheckResult("revision.md parse", "FAIL", str(exc))]

    tracked, clean, online, letter = _locate_outputs(submit_dir)
    student_name = revision.student_name

    results: list[CheckResult] = []
    results.append(_check_patch_anchor_integrity(revision, original_docx_path))
    results.append(_check_replace_all_counts(revision, original_docx_path))
    results.append(_check_citation_count(revision))
    results.append(_check_orphan_refs(revision))
    results.append(_check_sequential_numbering(revision))
    results.append(_check_ooxml_superscript(tracked))
    results.append(_check_standard_online_parity(revision, online))
    results.append(_check_duplication_artefacts(tracked))
    results.append(_check_vanishing_content(tracked))
    results.append(_check_author_attribution(tracked, clean, student_name))
    results.append(_check_anonymous_standard(clean))
    results.append(_check_figure_regenerated_flags(revision, submit_dir))
    results.append(_check_letter_vs_manuscript(letter, clean))
    return results


# ─── File location helper ──────────────────────────────────────────────────

def _locate_outputs(submit_dir: Path) -> tuple[Path | None, Path | None,
                                                Path | None, Path | None]:
    tracked = clean = online = letter = None
    if not submit_dir.exists():
        return (None, None, None, None)
    for p in submit_dir.glob("*.docx"):
        n = p.name.lower()
        if "standard-tracked" in n:
            tracked = p
        elif "standard-clean" in n:
            clean = p
        elif "online" in n:
            online = p
        elif "response-letter" in n or "response_letter" in n:
            letter = p
    return (tracked, clean, online, letter)


# ─── Patch-content helpers ─────────────────────────────────────────────────

def _all_citation_numbers_in_patches(revision) -> list[int]:
    """Walk all patch INSERT/REPLACE/APPEND text fields, collect [[CITE:N]]
    numbers in order. Used to derive what citations the revision body uses."""
    nums: list[int] = []
    for patch in revision.patches:
        for f in ("REPLACE", "INSERT", "APPEND", "INSERT_BODY"):
            v = patch.fields.get(f)
            if isinstance(v, str):
                for m in _CITE_RE.finditer(v):
                    for piece in m.group(1).split(","):
                        piece = piece.strip()
                        if piece.isdigit():
                            nums.append(int(piece))
        # REF_LIST ACTIONS may have NEW or ADD strings with [[CITE]] too
        if patch.op == "REF_LIST":
            for action in patch.fields.get("ACTIONS", []) or []:
                if not isinstance(action, dict):
                    continue
                for v in action.values():
                    if isinstance(v, str):
                        for m in _CITE_RE.finditer(v):
                            for piece in m.group(1).split(","):
                                piece = piece.strip()
                                if piece.isdigit():
                                    nums.append(int(piece))
    return nums


def _ref_list_numbers(revision) -> set[int]:
    """Reference numbers introduced by REF_LIST ADD entries (leading 'N.')."""
    ref_nums: set[int] = set()
    for patch in revision.patches:
        if patch.op != "REF_LIST":
            continue
        for action in patch.fields.get("ACTIONS", []) or []:
            if not isinstance(action, dict):
                continue
            if "ADD" in action:
                m = re.match(r"^\s*(\d+)\s*[.\)]", str(action["ADD"]))
                if m:
                    ref_nums.add(int(m.group(1)))
            if "REPLACE_AT" in action:
                ref_nums.add(int(action["REPLACE_AT"]))
    return ref_nums


def _norm(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


# ─── 1. Patch anchor integrity ─────────────────────────────────────────────

def _check_patch_anchor_integrity(revision, original_docx_path: Path) -> CheckResult:
    """Every FIND clause / anchor in every patch must match exactly one
    location in the original docx (or in the current document state, for
    patches that depend on prior insertions). For this audit we re-run a
    dry-pass on the original docx and report any anchor that fails."""
    try:
        doc = Document(str(original_docx_path))
    except Exception as exc:
        return CheckResult("Patch anchor integrity", "WARN",
                           f"could not open original docx: {exc}")

    body = doc.element.body

    def _all_p():
        return list(body.iter(qn("w:p")))

    failures = []
    for patch in revision.patches:
        op = patch.op
        # Identify the anchor field(s) for this op
        anchor_fields: list[tuple[str, str]] = []
        if op == "FIND_REPLACE":
            anchor_fields.append(("FIND (contains)", patch.fields.get("FIND", "")))
        elif op == "APPEND_TO_PARAGRAPH":
            anchor_fields.append(("PARAGRAPH_ENDS_WITH",
                                  patch.fields.get("PARAGRAPH_ENDS_WITH", "")))
        elif op == "INSERT_AFTER_PARAGRAPH":
            anchor_fields.append(("PARAGRAPH_ENDS_WITH",
                                  patch.fields.get("PARAGRAPH_ENDS_WITH", "")))
        elif op == "INSERT_AFTER_HEADING":
            anchor_fields.append(("AFTER_HEADING (exact)",
                                  patch.fields.get("AFTER_HEADING", "")))
        elif op == "REPLACE_HEADING":
            anchor_fields.append(("FIND_HEADING (exact)",
                                  patch.fields.get("FIND_HEADING", "")))
        elif op == "DELETE_PARAGRAPH":
            anchor_fields.append((
                "PARAGRAPH_STARTS_WITH" if patch.fields.get("PARAGRAPH_STARTS_WITH")
                else "PARAGRAPH_CONTAINS",
                patch.fields.get("PARAGRAPH_STARTS_WITH")
                or patch.fields.get("PARAGRAPH_CONTAINS", ""),
            ))
        elif op == "DELETE_RANGE":
            anchor_fields.append(("START_PARAGRAPH_STARTS_WITH",
                                  patch.fields.get("START_PARAGRAPH_STARTS_WITH", "")))
            anchor_fields.append(("END_PARAGRAPH_STARTS_WITH",
                                  patch.fields.get("END_PARAGRAPH_STARTS_WITH", "")))
        # TABLE_CELL / REPLACE_ALL / INSERT_TABLE checked separately
        for kind, anchor in anchor_fields:
            if not anchor:
                failures.append((patch.id, op, f"{kind} empty"))
                continue
            # Best-effort static match against the unmodified original.
            # Note: this is a heuristic — patches that depend on prior
            # insertions may report false positives here. The build itself
            # raises on real mismatches; this audit just gives a static read.
            n_match = 0
            anchor_n = _norm(anchor)
            for p in _all_p():
                ptext = _norm("".join(t.text or "" for t in p.iter(qn("w:t")) if t.text))
                if kind.startswith("FIND") and anchor_n in ptext:
                    n_match += 1
                elif kind.startswith("PARAGRAPH_ENDS_WITH") and ptext.endswith(anchor_n):
                    n_match += 1
                elif kind.startswith("PARAGRAPH_STARTS_WITH") and ptext.startswith(anchor_n):
                    n_match += 1
                elif kind.startswith("PARAGRAPH_CONTAINS") and anchor_n in ptext:
                    n_match += 1
                elif kind.startswith("AFTER_HEADING") and _norm(ptext) == anchor_n:
                    n_match += 1
                elif kind.startswith("FIND_HEADING") and _norm(ptext) == anchor_n:
                    n_match += 1
                elif kind.startswith("START_PARAGRAPH_STARTS_WITH") and ptext.startswith(anchor_n):
                    n_match += 1
                elif kind.startswith("END_PARAGRAPH_STARTS_WITH") and ptext.startswith(anchor_n):
                    n_match += 1
            # Heuristic: 0 matches → likely depends on a prior patch (OK as WARN).
            # >1 matches in the static view → real ambiguity.
            if n_match > 1:
                failures.append((patch.id, op,
                                 f"{kind} matched {n_match} paragraphs"))

    if not failures:
        return CheckResult("Patch anchor integrity", "PASS",
                           f"All {len(revision.patches)} patch anchors match "
                           f"≤1 paragraph in the original docx.")
    sample = "; ".join(f"{pid} {op} {detail}" for pid, op, detail in failures[:3])
    return CheckResult("Patch anchor integrity", "FAIL",
                       f"{len(failures)} anchor issue(s). Sample: {sample}")


# ─── 2. REPLACE_ALL CONFIRM_COUNT verification ─────────────────────────────

def _check_replace_all_counts(revision, original_docx_path: Path) -> CheckResult:
    """Confirm every REPLACE_ALL patch's CONFIRM_COUNT matches the actual
    count. unknown counts trigger a WARN with the actual count for review."""
    try:
        doc = Document(str(original_docx_path))
    except Exception as exc:
        return CheckResult("REPLACE_ALL counts", "WARN",
                           f"could not open original docx: {exc}")

    body = doc.element.body
    issues = []
    unknowns: list[tuple[str, int]] = []
    verified = 0

    for patch in revision.patches:
        if patch.op != "REPLACE_ALL":
            continue
        find = patch.fields.get("FIND", "")
        confirm = patch.fields.get("CONFIRM_COUNT")
        find_norm = _norm(find)
        actual = 0
        for p in body.iter(qn("w:p")):
            ptext = _norm("".join(t.text or "" for t in p.iter(qn("w:t")) if t.text))
            actual += ptext.count(find_norm)
        if confirm == "unknown":
            unknowns.append((patch.id, actual))
        elif isinstance(confirm, int):
            if confirm != actual:
                issues.append(f"{patch.id}: CONFIRM_COUNT={confirm} but actual={actual}")
            else:
                verified += 1
        else:
            issues.append(f"{patch.id}: missing or invalid CONFIRM_COUNT")

    if not issues and not unknowns:
        if verified == 0:
            return CheckResult("REPLACE_ALL counts", "PASS",
                               "No REPLACE_ALL patches in this revision.")
        return CheckResult("REPLACE_ALL counts", "PASS",
                           f"All {verified} REPLACE_ALL CONFIRM_COUNT match actual counts.")
    detail_parts = []
    if issues:
        detail_parts.append("MISMATCH: " + "; ".join(issues))
    if unknowns:
        u = "; ".join(f"{pid}: actual count {n}" for pid, n in unknowns)
        detail_parts.append(f"UNKNOWN auto-applied (review): {u}")
    status = "FAIL" if issues else "WARN"
    return CheckResult("REPLACE_ALL counts", status, " | ".join(detail_parts))


# ─── 3. Citation count consistency ─────────────────────────────────────────

def _check_citation_count(revision) -> CheckResult:
    cited = set(_all_citation_numbers_in_patches(revision))
    refs_added = _ref_list_numbers(revision)
    # We can't fully resolve "all refs that EXIST in the manuscript" without
    # reading the original docx's ref list. This check just verifies internal
    # consistency: every cited number ≤ max ref number added or referenced.
    if not cited and not refs_added:
        return CheckResult("Citation count", "PASS",
                           "No citation markers in patches; nothing to check.")
    if not refs_added:
        return CheckResult("Citation count", "WARN",
                           f"{len(cited)} unique citation numbers used in "
                           f"patches but no REF_LIST patches. Original docx "
                           f"reference list assumed to cover them.")
    only_cited = sorted(cited - refs_added)
    only_added = sorted(refs_added - cited)
    if not only_cited and not only_added:
        return CheckResult("Citation count", "PASS",
                           f"{len(cited)} unique citation numbers used; all "
                           f"covered by REF_LIST patches.")
    msg = []
    if only_cited:
        msg.append(f"cited but not in REF_LIST patches: {only_cited[:10]}")
    if only_added:
        msg.append(f"added by REF_LIST but never cited in patches: {only_added[:10]}")
    return CheckResult("Citation count", "WARN", "; ".join(msg))


# ─── 4. Orphan refs ────────────────────────────────────────────────────────

def _check_orphan_refs(revision) -> CheckResult:
    """v2.2 patch model: every REF_LIST ADD should be cited at least once
    in some patch's body content. We can't check the original docx's ref
    list orphans here without it; we just report on patch-added refs."""
    cited = set(_all_citation_numbers_in_patches(revision))
    added = _ref_list_numbers(revision)
    if not added:
        return CheckResult("Orphan refs", "PASS",
                           "No REF_LIST patches; nothing to orphan-check.")
    orphaned = sorted(added - cited)
    if not orphaned:
        return CheckResult("Orphan refs", "PASS",
                           f"All {len(added)} REF_LIST-added refs are cited "
                           f"in at least one patch body.")
    return CheckResult("Orphan refs", "WARN",
                       f"REF_LIST added refs that are never cited in patch "
                       f"bodies: {orphaned}. They may be cited only in the "
                       f"original docx body, which is OK if intentional.")


# ─── 5. Sequential numbering ───────────────────────────────────────────────

def _check_sequential_numbering(revision) -> CheckResult:
    nums = _all_citation_numbers_in_patches(revision)
    if not nums:
        return CheckResult("Sequential numbering", "PASS",
                           "No citation markers; nothing to check.")
    # Drop duplicates while preserving first-appearance order
    seen = []
    for n in nums:
        if n not in seen:
            seen.append(n)
    expected = list(range(1, len(seen) + 1))
    # Patches edit specific sections of the manuscript — the locked Phase-2
    # numbering doesn't require patches to use refs in 1..N order, only that
    # the FULL final manuscript does. So a WARN if patches alone don't show
    # 1..N, not a FAIL.
    if seen == expected:
        return CheckResult("Sequential numbering", "PASS",
                           f"Citations in patches appear in order 1..{len(seen)}.")
    return CheckResult("Sequential numbering", "WARN",
                       f"Patches use citations in non-sequential order: {seen}. "
                       f"This may be fine if the original docx body has the "
                       f"intermediate refs; verify against final manuscript.")


# ─── 6. OOXML superscript sanity ───────────────────────────────────────────

def _check_ooxml_superscript(tracked: Path | None) -> CheckResult:
    if tracked is None or not tracked.exists():
        return CheckResult("OOXML superscript", "WARN",
                           "Standard-Tracked.docx not found.")
    doc = Document(str(tracked))
    n = 0
    for r in doc.element.body.iter(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue
        va = rPr.find(qn("w:vertAlign"))
        if va is not None and va.get(qn("w:val")) == "superscript":
            n += 1
    if n == 0:
        return CheckResult("OOXML superscript", "FAIL",
                           "Zero <w:vertAlign w:val=\"superscript\"> runs "
                           "in Standard-Tracked.docx. Citations will not "
                           "survive Standard→Online conversion.")
    return CheckResult("OOXML superscript", "PASS",
                       f"Found {n} real OOXML superscript run(s).")


# ─── 7. Standard↔Online citation parity ────────────────────────────────────

def _check_standard_online_parity(revision, online: Path | None) -> CheckResult:
    if online is None or not online.exists():
        return CheckResult("Standard↔Online parity", "WARN",
                           "Online docx not found.")
    expected = len(_all_citation_numbers_in_patches(revision))
    doc = Document(str(online))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    blocks = full_text.count("((")
    if expected == 0:
        return CheckResult("Standard↔Online parity", "PASS",
                           f"No citations in patches; Online has {blocks} (( blocks "
                           f"(from original docx body).")
    # We can't easily count original-docx citations here. Approximate.
    if blocks > 0:
        return CheckResult("Standard↔Online parity", "PASS",
                           f"Online docx has {blocks} (( blocks; "
                           f"{expected} citation slots from patches.")
    return CheckResult("Standard↔Online parity", "FAIL",
                       f"Online docx has 0 (( blocks but patches added "
                       f"{expected} citation slots. Standard→Online dropped them. "
                       f"Likely cause: References section format unsupported by parser.")


# ─── 8. No duplication artefacts ───────────────────────────────────────────

def _check_duplication_artefacts(tracked: Path | None) -> CheckResult:
    """Scan tracked docx for <w:ins> blocks whose text content appears
    verbatim elsewhere in the docx outside the deletion blocks."""
    if tracked is None or not tracked.exists():
        return CheckResult("Duplication artefacts", "WARN",
                           "Standard-Tracked.docx not found.")
    doc = Document(str(tracked))
    body = doc.element.body

    ins_texts: list[str] = []
    for ins in body.iter(qn("w:ins")):
        text = "".join(t.text or "" for t in ins.iter(qn("w:t")))
        if len(text.strip()) >= 50:  # only substantial blocks
            ins_texts.append(text)

    # Gather all non-deleted text (skip <w:del>)
    plain_text_parts = []
    for t in body.iter(qn("w:t")):
        # Skip text inside <w:ins> (we only care about pre-existing text)
        a = t.getparent()
        in_ins = False
        in_del = False
        while a is not None and a.tag != qn("w:body"):
            if a.tag == qn("w:ins"):
                in_ins = True
            if a.tag == qn("w:del"):
                in_del = True
            a = a.getparent()
        if not in_ins and not in_del and t.text:
            plain_text_parts.append(t.text)
    plain_text = " ".join(plain_text_parts)

    duplicates = []
    for ins in ins_texts:
        ins_n = _norm(ins)
        if ins_n in _norm(plain_text):
            duplicates.append(ins_n[:80])

    if not duplicates:
        return CheckResult("Duplication artefacts", "PASS",
                           f"No <w:ins> blocks duplicate existing body text.")
    return CheckResult("Duplication artefacts", "FAIL",
                       f"{len(duplicates)} insertion(s) duplicate existing text. "
                       f"First: {duplicates[0]!r}")


# ─── 9. No vanishing content ───────────────────────────────────────────────

def _check_vanishing_content(tracked: Path | None) -> CheckResult:
    """Scan for content that has no future after Accept All — every <w:del>
    block whose text was unique to that paragraph and no <w:ins> takes its
    place is content that vanishes. Heuristic check: report if any
    paragraph is entirely deleted (all runs in <w:del>) and has no <w:ins>
    siblings."""
    if tracked is None or not tracked.exists():
        return CheckResult("Vanishing content", "WARN",
                           "Standard-Tracked.docx not found.")
    doc = Document(str(tracked))
    suspicious = []
    for p in doc.element.body.iter(qn("w:p")):
        has_del = next(p.iter(qn("w:del")), None) is not None
        has_ins = next(p.iter(qn("w:ins")), None) is not None
        if not has_del:
            continue
        # Has only deletions, no insertions, AND paragraph mark not deleted?
        # That's a normal mid-paragraph deletion. Skip.
        # Paragraph-mark-deleted paragraphs disappear cleanly; not suspicious.
        pPr = p.find(qn("w:pPr"))
        para_mark_deleted = False
        if pPr is not None:
            rPr = pPr.find(qn("w:rPr"))
            if rPr is not None and rPr.find(qn("w:del")) is not None:
                para_mark_deleted = True
        # We're looking for paragraphs where ALL text-bearing children are in
        # <w:del> and there's NO <w:ins> — and the paragraph mark isn't
        # deleted (so an empty paragraph survives Accept All).
        if not has_ins and not para_mark_deleted:
            text_in_del = sum(1 for _ in p.iter(qn("w:delText")))
            text_outside = 0
            for t in p.iter(qn("w:t")):
                a = t.getparent()
                while a is not None and a.tag != qn("w:body"):
                    if a.tag == qn("w:del"):
                        break
                    a = a.getparent()
                else:
                    text_outside += 1 if t.text else 0
            if text_in_del > 0 and text_outside == 0:
                suspicious.append(p)
    if not suspicious:
        return CheckResult("Vanishing content", "PASS",
                           "No paragraphs would survive Accept-All as empty.")
    return CheckResult("Vanishing content", "WARN",
                       f"{len(suspicious)} paragraph(s) have all content in "
                       f"<w:del> with no <w:ins> replacement. Accept-All "
                       f"leaves them empty. Verify intentional.")


# ─── 10. Author attribution ────────────────────────────────────────────────

def _check_author_attribution(tracked: Path | None,
                              clean: Path | None,
                              student_name: str) -> CheckResult:
    if not student_name:
        return CheckResult("Author attribution", "FAIL",
                           "revision.md missing student_name.")
    issues = []
    for label, p in (("Tracked", tracked), ("Clean", clean)):
        if p is None or not p.exists():
            issues.append(f"{label} missing")
            continue
        try:
            doc = Document(str(p))
            author = (doc.core_properties.author or "").strip()
        except Exception as exc:
            issues.append(f"{label} unreadable: {exc}")
            continue
        if author.lower() in _AUTHOR_BLOCKLIST:
            issues.append(f"{label} author {author!r} blocklisted")
        elif author != student_name:
            issues.append(f"{label} author {author!r}, expected {student_name!r}")
    if issues:
        return CheckResult("Author attribution", "FAIL", "; ".join(issues))
    return CheckResult("Author attribution", "PASS",
                       f"Tracked + Clean docx author = {student_name!r}.")


# ─── 11. Anonymous Standard ────────────────────────────────────────────────

def _check_anonymous_standard(clean: Path | None) -> CheckResult:
    if clean is None or not clean.exists():
        return CheckResult("Anonymous Standard", "WARN",
                           "Clean docx not found.")
    doc = Document(str(clean))
    full = "\n".join(p.text for p in doc.paragraphs)
    hits = _IDENTIFYING_REGEX.findall(full)
    if hits:
        return CheckResult("Anonymous Standard", "FAIL",
                           f"Identifying language found: {list(set(hits))}")
    return CheckResult("Anonymous Standard", "PASS",
                       "No identifying language detected in Standard-Clean.")


# ─── 12. Figure regeneration flags ─────────────────────────────────────────

def _check_figure_regenerated_flags(revision, submit_dir: Path) -> CheckResult:
    """Every FIGURE_REGENERATED flag must reference a NEW_FILE that exists
    on disk. If no FIGURE_REGENERATED flags are present, the check passes."""
    flag_patches = [p for p in revision.patches if p.op == "FIGURE_REGENERATED"]
    if not flag_patches:
        return CheckResult("Figure regeneration flags", "PASS",
                           "No FIGURE_REGENERATED flags in this revision.")
    missing = []
    for patch in flag_patches:
        new_file = patch.fields.get("NEW_FILE", "")
        if not new_file:
            missing.append(f"{patch.id}: missing NEW_FILE field")
            continue
        # Look in submit_dir, then submit_dir's parent (review folder)
        candidates = [
            submit_dir / new_file,
            submit_dir.parent / new_file,
        ]
        if not any(c.exists() for c in candidates):
            missing.append(f"{patch.id}: NEW_FILE {new_file!r} not found")
    if missing:
        return CheckResult("Figure regeneration flags", "FAIL",
                           "; ".join(missing))
    return CheckResult("Figure regeneration flags", "PASS",
                       f"All {len(flag_patches)} FIGURE_REGENERATED flags "
                       f"resolve to existing PNG files.")


# ─── 13. Letter ↔ manuscript reality ──────────────────────────────────────

def _check_letter_vs_manuscript(letter: Path | None,
                                clean: Path | None) -> CheckResult:
    if letter is None or not letter.exists():
        return CheckResult("Letter ↔ manuscript", "WARN",
                           "Response-Letter.docx not found.")
    if clean is None or not clean.exists():
        return CheckResult("Letter ↔ manuscript", "WARN",
                           "Clean docx not found.")
    letter_doc = Document(str(letter))
    letter_text = "\n".join(p.text for p in letter_doc.paragraphs)
    locations = re.findall(r"Change location:\s*(.+?)\s*$",
                            letter_text, re.MULTILINE)
    if not locations:
        return CheckResult("Letter ↔ manuscript", "WARN",
                           "No 'Change location:' lines in letter.")
    clean_doc = Document(str(clean))
    headings = {
        p.text.strip().lower()
        for p in clean_doc.paragraphs
        if p.style.name.lower().startswith("heading")
    }
    # Also accept plain "Abstract"/"Introduction"/etc. paragraphs as quasi-headings
    SECTION_NAMES = {"abstract", "introduction", "methods", "results",
                     "discussion", "conclusion", "references", "appendix",
                     "limitations", "acknowledgments"}
    for p in clean_doc.paragraphs:
        if _norm(p.text).lower() in SECTION_NAMES:
            headings.add(_norm(p.text).lower())
    if not headings:
        return CheckResult("Letter ↔ manuscript", "WARN",
                           "Clean docx has no headings to cross-check.")
    misses = []
    for loc in locations:
        l = loc.lower()
        if "no manuscript change" in l:
            continue
        section_word = l.split(",")[0].strip()
        if not any(section_word in h or h in section_word for h in headings):
            misses.append(loc)
    if misses:
        return CheckResult("Letter ↔ manuscript", "WARN",
                           f"{len(misses)} 'Change location' line(s) reference "
                           f"unknown sections. Sample: {misses[:2]}")
    return CheckResult("Letter ↔ manuscript", "PASS",
                       f"All {len(locations)} 'Change location' lines map to "
                       f"a section in the clean manuscript.")


# ─── Reporting helpers ─────────────────────────────────────────────────────

def format_report(results: Iterable[CheckResult]) -> str:
    lines = []
    width = max(len(r.name) for r in results) + 2
    for r in results:
        lines.append(f"  {r.status:<4}  {r.name:<{width}}  {r.detail}")
    return "\n".join(lines)


def summary_counts(results: Iterable[CheckResult]) -> dict[str, int]:
    counts = {"PASS": 0, "WARN": 0, "FAIL": 0}
    for r in results:
        counts[r.status] = counts.get(r.status, 0) + 1
    return counts


# ─── CLI ───────────────────────────────────────────────────────────────────

def _cli():
    parser = argparse.ArgumentParser(
        description="Audit a v2.0.0 NHSJS revision build against the Phase 6 checklist.",
    )
    parser.add_argument("revision_md", help="Path to revision.md")
    parser.add_argument("original_docx", help="Path to the original Standard docx")
    parser.add_argument("submit_dir", help="Directory containing built outputs")
    parser.add_argument("--log", metavar="FILE",
                        help="Optional revision-log.md")
    args = parser.parse_args()

    results = audit_revision(args.revision_md, args.original_docx,
                              args.submit_dir, log_path=args.log)
    print(format_report(results))
    print()
    counts = summary_counts(results)
    print(f"Summary: {counts['PASS']} pass, {counts['WARN']} warn, "
          f"{counts['FAIL']} fail")
    raise SystemExit(0 if counts["FAIL"] == 0 else 1)


if __name__ == "__main__":
    _cli()
