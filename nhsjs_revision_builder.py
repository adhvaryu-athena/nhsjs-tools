"""
nhsjs_revision_builder.py — v2.0.0 patch-application engine
============================================================
Consumes a v2.2 `revision.md` (sequence of anchored patches) and an
original Standard manuscript .docx, applies patches in spec-defined
order with FIND-must-match-verbatim verification, emits 5 NHSJS
submission outputs.

This is a full rewrite from v1.0.0's marker-in-full-manuscript model.
v1.0.0 wrote a fresh docx from a marker-laden manuscript; v2.0.0
applies anchored patches to the live XML of the original docx. The
anti-hallucination argument: the bot never re-emits unchanged text, so
multi-thousand-word manuscripts don't trigger transcription drift.

Public API
----------
    parse_revision(revision_md_path) -> RevisionDoc
        Read, parse, and validate revision.md. Raises BuilderError on
        malformed input.

    build_all(revision_md_path, original_docx_path,
              revision_log_md_path, output_dir) -> dict
        Run the full pipeline. Produces 5 outputs + FILES.md.
        Returns a dict mapping output kind to absolute path.

    apply_patches(revision, original_docx_path, output_path) -> ApplyReport
        Lower-level: apply patches against original_docx_path, save
        Standard-Tracked.docx. Returns per-patch success/failure report.

Patch vocabulary (per v2.2 spec)
--------------------------------
    Text:       FIND_REPLACE, APPEND_TO_PARAGRAPH, INSERT_AFTER_PARAGRAPH,
                INSERT_AFTER_HEADING, REPLACE_HEADING, DELETE_PARAGRAPH,
                DELETE_RANGE
    References: REF_LIST  (actions: ADD, REPLACE_AT, RENUMBER)
    Tables:     TABLE_CELL, INSERT_TABLE
    Wide:       REPLACE_ALL  (with CONFIRM_COUNT integer or 'unknown')
    Flag:       FIGURE_REGENERATED  (audit-only, no docx change)

Citation marker `[[CITE:N,N,N]]` resolves to real OOXML superscript runs.
Unit superscripts (km², m³, CO₂, cm⁻¹) stay as literal Unicode in patch
text and are preserved as-is.

REF_LIST semantics (locked v2.0.0)
----------------------------------
    ADD          — parses leading "N." from the entry string. Inserts
                   the paragraph at position N in the reference list
                   (between current N-1 and N+1).
    REPLACE_AT N — swaps the entry at position N in place. NEW carries
                   its own number prefix.
    RENUMBER     — takes an explicit MAP: {old: new}. Engine applies
                   the map to every [[CITE:N]] superscript run in the
                   body (after all text patches are applied).

Body [[CITE:N]] markers in any patch use Phase-2-locked final numbers
unless a RENUMBER patch is present to remap them.

Order of operations (Phase 4 spec)
----------------------------------
    1. Text patches in document order
    2. Table patches
    3. Scope ops (REPLACE_ALL)
    4. Reference operations (REF_LIST) — last
    5. Resolve [[CITE:N]] to real <w:vertAlign w:val="superscript">
    6. Set core_properties.author = student_name
    7. Save Standard-Tracked.docx
    8. Accept all <w:ins>, drop all <w:del>, save Standard-Clean.docx
    9. Pipe Standard-Clean through Standard→Online → save Online.docx
   10. Generate Response-Letter.docx from revision-log.md
   11. Generate Student-Handoff.md from revision-log.md + patch summary

Strict-mode failures
--------------------
    The build raises BuilderError immediately if any:
    - Missing/malformed YAML frontmatter
    - Patch anchor matches zero or more than one location
    - REPLACE_ALL CONFIRM_COUNT integer doesn't match actual occurrences
      (CONFIRM_COUNT: unknown auto-applies; flagged in audit)
    - Unresolved [[CITE:REF_NEEDED_*]] placeholder
    - Unicode superscript glyphs ¹²³ used in citation slots (use [[CITE:N]])
"""
from __future__ import annotations

import argparse
import copy
import datetime
import json
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable

import yaml
from docx import Document
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_VERSION = "2.0.0"


# ─── Constants ─────────────────────────────────────────────────────────────

RED = RGBColor(0xC0, 0x00, 0x00)

_CITE_RE = re.compile(r"\[\[CITE:([^\]]+)\]\]")
_REF_NEEDED_RE = re.compile(r"\[\[CITE:[^\]]*REF_NEEDED[^\]]*\]\]", re.IGNORECASE)
_UNICODE_SUPERSCRIPT_RE = re.compile(r"[²³¹⁰-⁹]")
_VALID_CITE_BODY_RE = re.compile(r"^\s*\d+(\s*,\s*\d+)*\s*$")
_REF_NUMBER_PREFIX_RE = re.compile(r"^\s*(\d+)\s*[.\)]\s+")

# Unit superscripts that are NOT citations — preserve as-is
_UNIT_PRECEDES_RE = re.compile(r"(?:km|cm|mm|m|ha|kg|g|s|CO|H|N|O|kJ)$")


# ─── Errors ────────────────────────────────────────────────────────────────

class BuilderError(Exception):
    """Raised on any strict-mode failure during build."""

    def __init__(self, message: str, *, kind: str = "io",
                 patch_id: str | None = None,
                 line: int | None = None):
        super().__init__(message)
        self.kind = kind
        self.patch_id = patch_id
        self.line = line


# ─── Patch + RevisionDoc data classes ──────────────────────────────────────

@dataclass
class Patch:
    """One v2.2 patch block."""
    id: str
    op: str
    fields: dict[str, Any] = field(default_factory=dict)


@dataclass
class RevisionDoc:
    """Parsed revision.md."""
    student_name: str
    manuscript_title: str
    reviewer_decision_date: str
    target_submit_date: str
    frontmatter: dict[str, Any]
    patches: list[Patch]
    source_path: str | None = None


@dataclass
class ApplyResult:
    """One patch's application outcome."""
    patch_id: str
    op: str
    status: str  # "applied" | "failed"
    detail: str = ""
    match_count: int | None = None  # for REPLACE_ALL


@dataclass
class ApplyReport:
    results: list[ApplyResult] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def applied(self) -> int:
        return sum(1 for r in self.results if r.status == "applied")

    @property
    def failed(self) -> int:
        return sum(1 for r in self.results if r.status == "failed")


# ─── Whitespace normalization ──────────────────────────────────────────────

_NBSP = " "
_ZWJ = "‍"


def _norm(s: str) -> str:
    """v2.2 anchor matching norm: tolerant of whitespace, strict on
    everything else. Collapse \\s+ → ' '. NBSP + ZWJ treated as space.
    Strip outer."""
    if s is None:
        return ""
    s = s.replace(_NBSP, " ").replace(_ZWJ, " ")
    return re.sub(r"\s+", " ", s).strip()


# ─── Patch parsing ─────────────────────────────────────────────────────────

def parse_revision(revision_md_path: str) -> RevisionDoc:
    """Read revision.md, parse YAML frontmatter + patch blocks, validate.

    Raises BuilderError on any malformed input that would cause the build
    to produce wrong output.
    """
    path = Path(revision_md_path)
    if not path.exists():
        raise BuilderError(f"revision.md not found: {revision_md_path}", kind="io")

    raw = path.read_text(encoding="utf-8")

    # 1. Frontmatter
    m = re.match(r"^---\s*\n(.*?)\n---\s*\n(.*)", raw, re.DOTALL)
    if not m:
        raise BuilderError(
            "revision.md missing YAML frontmatter (`--- ... ---` block).",
            kind="frontmatter",
        )
    try:
        fm = yaml.safe_load(m.group(1)) or {}
    except yaml.YAMLError as exc:
        raise BuilderError(
            f"revision.md frontmatter is not valid YAML: {exc}",
            kind="frontmatter",
        ) from exc
    body = m.group(2)

    if not fm.get("student_name"):
        raise BuilderError(
            "revision.md frontmatter missing required `student_name`.",
            kind="frontmatter",
        )
    if not fm.get("manuscript_title"):
        raise BuilderError(
            "revision.md frontmatter missing required `manuscript_title`.",
            kind="frontmatter",
        )

    # 2. Body-level strict checks (REF_NEEDED, Unicode superscript)
    body_line_offset = raw.count("\n", 0, m.start(2))

    rn = _REF_NEEDED_RE.search(body)
    if rn:
        line = body.count("\n", 0, rn.start()) + 1 + body_line_offset
        raise BuilderError(
            f"Unresolved citation placeholder {rn.group(0)} at line {line}. "
            "Run Reference Hunter to fill before building.",
            kind="ref_needed",
            line=line,
        )

    us = _UNICODE_SUPERSCRIPT_RE.search(body)
    if us:
        # Unit superscripts are OK in patch text per v2.2 — the spec says
        # [[CITE:N]] is for citations only. If the bot wrote ¹²³ in a
        # citation slot, that's an error. Heuristic: if preceded by a unit,
        # it's fine; if preceded by a word, it's likely a citation gone wrong.
        idx = us.start()
        preceding = body[max(0, idx - 20):idx]
        if _UNIT_PRECEDES_RE.search(preceding):
            pass  # unit superscript, OK
        else:
            line = body.count("\n", 0, idx) + 1 + body_line_offset
            raise BuilderError(
                f"Unicode superscript glyph {us.group(0)!r} at line {line}. "
                "If this is a citation, replace with [[CITE:N]] marker. "
                "If it's a unit (e.g., km²) preceded by a recognized unit "
                "string, this check should not have fired — file a bug.",
                kind="unicode_superscript",
                line=line,
            )

    # 3. Parse patch blocks
    patches: list[Patch] = []
    blocks = re.split(r"\n(?=## P\d)", body)
    for block in blocks:
        block = block.strip()
        if not block.startswith("## P"):
            continue
        patches.append(_parse_patch_block(block, body_line_offset))

    return RevisionDoc(
        student_name=str(fm["student_name"]),
        manuscript_title=str(fm["manuscript_title"]),
        reviewer_decision_date=str(fm.get("reviewer_decision_date", "")),
        target_submit_date=str(fm.get("target_submit_date", "")),
        frontmatter=fm,
        patches=patches,
        source_path=str(path),
    )


def _parse_patch_block(block: str, body_line_offset: int) -> Patch:
    """Parse a single `## P{N}` block into a Patch."""
    lines = block.splitlines()
    header = lines[0]
    pid_m = re.match(r"^##\s+(P\d+)", header)
    if not pid_m:
        raise BuilderError(
            f"Patch header malformed: {header!r}",
            kind="parse",
        )
    pid = pid_m.group(1)

    fields: dict[str, Any] = {}
    current_list_key: str | None = None
    current_list: list[Any] | None = None

    i = 1
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        # Try `KEY: value` (uppercase + underscores allowed)
        km = re.match(r"^([A-Z][A-Z0-9_]*):\s*(.*)$", line)
        if km:
            current_list_key = None
            current_list = None
            key, val = km.group(1), km.group(2)
            if not val.strip():
                # Possibly a list start (next lines `  - ...`)
                current_list_key = key
                current_list = []
                fields[key] = current_list
            else:
                fields[key] = _coerce_value(val)
            i += 1
            continue
        # List item under current key
        if current_list is not None and re.match(r"^\s*-\s", line):
            # Could be an inline scalar (`- ADD: "..."`) or a structured
            # sub-block over multiple lines. For v2.0 we support the common
            # REF_LIST shape: each `- ACTION: value` is a dict with that
            # action as a single key.
            item_text = line.strip()[1:].strip()  # strip leading `-`
            mm = re.match(r"^([A-Z][A-Z0-9_]*):\s*(.*)$", item_text)
            if mm:
                action_key = mm.group(1)
                action_val = mm.group(2)
                if not action_val.strip():
                    # Multi-line sub-block following
                    sub: dict[str, Any] = {action_key: None}
                    j = i + 1
                    while j < len(lines):
                        sub_line = lines[j]
                        if not sub_line.strip():
                            j += 1
                            continue
                        if re.match(r"^\s*-\s|^[A-Z][A-Z0-9_]*:", sub_line):
                            break
                        sm = re.match(r"^\s+([A-Z][A-Z0-9_]*):\s*(.*)$", sub_line)
                        if sm:
                            sub[sm.group(1)] = _coerce_value(sm.group(2))
                        j += 1
                    current_list.append(sub)
                    i = j
                    continue
                else:
                    current_list.append({action_key: _coerce_value(action_val)})
                    i += 1
                    continue
            else:
                current_list.append(_coerce_value(item_text))
                i += 1
                continue
        i += 1

    op = fields.pop("OP", None)
    if not op:
        raise BuilderError(
            f"Patch {pid} missing `OP:` field.",
            kind="parse",
            patch_id=pid,
        )

    return Patch(id=pid, op=str(op), fields=fields)


def _coerce_value(s: str) -> Any:
    s = s.strip()
    # JSON-style quoted string
    if s.startswith('"') and s.endswith('"'):
        try:
            return json.loads(s)
        except json.JSONDecodeError:
            return s[1:-1]
    # YAML-style dict (e.g. RENUMBER MAP)
    if s.startswith("{") and s.endswith("}"):
        try:
            return yaml.safe_load(s)
        except Exception:
            return s
    # Integer
    if re.fullmatch(r"-?\d+", s):
        return int(s)
    # 'unknown' keyword (REPLACE_ALL)
    if s == "unknown":
        return "unknown"
    return s


# ─── OOXML helpers ─────────────────────────────────────────────────────────

def _author_attrs(student_name: str) -> dict[str, str]:
    """Common attributes for w:ins / w:del elements."""
    now = datetime.datetime.now().replace(microsecond=0).isoformat() + "Z"
    return {
        qn("w:id"): "0",  # caller assigns
        qn("w:author"): student_name,
        qn("w:date"): now,
    }


def _next_change_id() -> str:
    """Monotonic IDs for w:ins / w:del so Word can pair them."""
    _next_change_id.counter += 1  # type: ignore[attr-defined]
    return str(_next_change_id.counter)


_next_change_id.counter = 0  # type: ignore[attr-defined]


def _make_run(text: str, *, superscript: bool = False) -> OxmlElement:
    """Build a <w:r> with the given text. If superscript=True, wraps in
    <w:vertAlign w:val="superscript">."""
    r = OxmlElement("w:r")
    if superscript:
        rPr = OxmlElement("w:rPr")
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rPr.append(va)
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _make_runs_for_text(text: str) -> list[OxmlElement]:
    """Convert plain text + [[CITE:N]] markers into a list of <w:r> elements,
    with citations emitted as real superscript runs."""
    runs: list[OxmlElement] = []
    i = 0
    for m in _CITE_RE.finditer(text):
        if m.start() > i:
            runs.append(_make_run(text[i:m.start()]))
        body = m.group(1)
        if not _VALID_CITE_BODY_RE.match(body):
            raise BuilderError(
                f"Malformed citation marker [[CITE:{body}]]. "
                "Content must be comma-separated digits.",
                kind="malformed_cite",
            )
        digits = re.sub(r"\s+", "", body)
        runs.append(_make_run(digits, superscript=True))
        i = m.end()
    if i < len(text):
        runs.append(_make_run(text[i:]))
    return runs


def _wrap_ins(runs: list[OxmlElement], student_name: str) -> OxmlElement:
    """Wrap runs in a <w:ins> element."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), _next_change_id())
    ins.set(qn("w:author"), student_name)
    ins.set(qn("w:date"), datetime.datetime.now().replace(microsecond=0).isoformat() + "Z")
    for r in runs:
        ins.append(r)
    return ins


def _wrap_del(runs: list[OxmlElement], student_name: str) -> OxmlElement:
    """Wrap runs in <w:del>, converting <w:t> → <w:delText>."""
    d = OxmlElement("w:del")
    d.set(qn("w:id"), _next_change_id())
    d.set(qn("w:author"), student_name)
    d.set(qn("w:date"), datetime.datetime.now().replace(microsecond=0).isoformat() + "Z")
    for r in runs:
        # Convert <w:t> to <w:delText>
        for t in r.findall(qn("w:t")):
            new_t = OxmlElement("w:delText")
            new_t.text = t.text
            for k, v in t.attrib.items():
                new_t.set(k, v)
            r.replace(t, new_t)
        d.append(r)
    return d


# ─── Anchor finding ────────────────────────────────────────────────────────

def _paragraph_text(p_elem) -> str:
    """All visible text in a paragraph including current insertions, excluding
    current deletions. (For matching against patches at apply time, anchors
    should still match the pre-patch text — but since we apply patches in
    order, anchors after an earlier patch may need to see post-patch state.
    We match against text-as-now: insertions count, deletions don't.)"""
    parts = []
    for t in p_elem.iter(qn("w:t")):
        # Skip if ancestor is <w:del>
        a = t.getparent()
        in_del = False
        while a is not None and a.tag != qn("w:body"):
            if a.tag == qn("w:del"):
                in_del = True
                break
            a = a.getparent()
        if not in_del:
            parts.append(t.text or "")
    return "".join(parts)


def _paragraph_style(p_elem) -> str | None:
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return None
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        return None
    return pStyle.get(qn("w:val"))


def _find_paragraphs_containing(doc_body, needle: str) -> list:
    """Return all <w:p> elements whose text contains `needle` (normalized)."""
    needle_norm = _norm(needle)
    if not needle_norm:
        return []
    matches = []
    for p in doc_body.iter(qn("w:p")):
        if needle_norm in _norm(_paragraph_text(p)):
            matches.append(p)
    return matches


def _find_paragraphs_starts_with(doc_body, needle: str) -> list:
    needle_norm = _norm(needle)
    if not needle_norm:
        return []
    return [p for p in doc_body.iter(qn("w:p"))
            if _norm(_paragraph_text(p)).startswith(needle_norm)]


def _find_paragraphs_ends_with(doc_body, needle: str) -> list:
    needle_norm = _norm(needle)
    if not needle_norm:
        return []
    return [p for p in doc_body.iter(qn("w:p"))
            if _norm(_paragraph_text(p)).endswith(needle_norm)]


def _expect_one(patch_id: str, op: str, kind: str, anchor: str,
                matches: list) -> Any:
    """v2.2 spec: anchor must match exactly one. Fail loudly otherwise."""
    if not matches:
        raise BuilderError(
            f"{op} anchor not found. Looked for {kind}: {anchor[:80]!r}",
            kind="anchor_missing",
            patch_id=patch_id,
        )
    if len(matches) > 1:
        raise BuilderError(
            f"{op} anchor matched {len(matches)} paragraphs (expected 1). "
            f"Anchor: {anchor[:80]!r}. Refine the anchor or use a more "
            f"specific patch.",
            kind="anchor_ambiguous",
            patch_id=patch_id,
        )
    return matches[0]


# ─── Op handlers ───────────────────────────────────────────────────────────

def _op_find_replace(doc, patch: Patch, student_name: str) -> ApplyResult:
    find = patch.fields.get("FIND", "")
    replace = patch.fields.get("REPLACE", "")
    if not find:
        raise BuilderError("FIND_REPLACE missing FIND", kind="parse",
                           patch_id=patch.id)
    matches = _find_paragraphs_containing(doc.element.body, find)
    p = _expect_one(patch.id, "FIND_REPLACE", "paragraph containing",
                    find, matches)
    _replace_text_in_paragraph(p, find, replace, student_name)
    return ApplyResult(patch_id=patch.id, op="FIND_REPLACE", status="applied")


def _op_append_to_paragraph(doc, patch: Patch, student_name: str) -> ApplyResult:
    anchor = patch.fields.get("PARAGRAPH_ENDS_WITH", "")
    append = patch.fields.get("APPEND", "")
    if not anchor or not append:
        raise BuilderError("APPEND_TO_PARAGRAPH missing PARAGRAPH_ENDS_WITH "
                           "or APPEND", kind="parse", patch_id=patch.id)
    matches = _find_paragraphs_ends_with(doc.element.body, anchor)
    p = _expect_one(patch.id, "APPEND_TO_PARAGRAPH", "paragraph ending with",
                    anchor, matches)
    # Append " {APPEND}" as an inserted run (with leading space if needed)
    text = " " + append if not append.startswith(" ") else append
    new_runs = _make_runs_for_text(text)
    ins = _wrap_ins(new_runs, student_name)
    p.append(ins)
    return ApplyResult(patch_id=patch.id, op="APPEND_TO_PARAGRAPH", status="applied")


def _op_insert_after_paragraph(doc, patch: Patch, student_name: str) -> ApplyResult:
    anchor = patch.fields.get("PARAGRAPH_ENDS_WITH", "")
    insert = patch.fields.get("INSERT", "")
    if not anchor or not insert:
        raise BuilderError("INSERT_AFTER_PARAGRAPH missing PARAGRAPH_ENDS_WITH "
                           "or INSERT", kind="parse", patch_id=patch.id)
    matches = _find_paragraphs_ends_with(doc.element.body, anchor)
    p = _expect_one(patch.id, "INSERT_AFTER_PARAGRAPH",
                    "paragraph ending with", anchor, matches)
    _insert_paragraph_after(p, insert, student_name)
    return ApplyResult(patch_id=patch.id, op="INSERT_AFTER_PARAGRAPH",
                       status="applied")


def _op_insert_after_heading(doc, patch: Patch, student_name: str) -> ApplyResult:
    heading = patch.fields.get("AFTER_HEADING", "")
    if not heading:
        raise BuilderError("INSERT_AFTER_HEADING missing AFTER_HEADING",
                           kind="parse", patch_id=patch.id)
    # Find paragraph whose text equals heading AND style is a Heading
    matches = []
    for p in doc.element.body.iter(qn("w:p")):
        text = _paragraph_text(p).strip()
        if _norm(text) == _norm(heading):
            matches.append(p)
    p = _expect_one(patch.id, "INSERT_AFTER_HEADING", "heading", heading, matches)
    # Optional: insert a new heading first, then body
    new_heading_text = patch.fields.get("INSERT_HEADING")
    if new_heading_text:
        level = int(patch.fields.get("HEADING_LEVEL", 2))
        _insert_heading_after(p, new_heading_text, level, student_name)
    body_text = patch.fields.get("INSERT_BODY") or patch.fields.get("INSERT", "")
    if body_text:
        _insert_paragraph_after(p, body_text, student_name)
    return ApplyResult(patch_id=patch.id, op="INSERT_AFTER_HEADING",
                       status="applied")


def _op_replace_heading(doc, patch: Patch, student_name: str) -> ApplyResult:
    find = patch.fields.get("FIND_HEADING", "")
    replace = patch.fields.get("REPLACE_HEADING", "")
    if not find:
        raise BuilderError("REPLACE_HEADING missing FIND_HEADING",
                           kind="parse", patch_id=patch.id)
    matches = []
    for p in doc.element.body.iter(qn("w:p")):
        if _norm(_paragraph_text(p)) == _norm(find):
            matches.append(p)
    p = _expect_one(patch.id, "REPLACE_HEADING", "heading text",
                    find, matches)
    _replace_text_in_paragraph(p, find, replace, student_name)
    return ApplyResult(patch_id=patch.id, op="REPLACE_HEADING", status="applied")


def _op_delete_paragraph(doc, patch: Patch, student_name: str) -> ApplyResult:
    starts_with = patch.fields.get("PARAGRAPH_STARTS_WITH")
    contains = patch.fields.get("PARAGRAPH_CONTAINS")
    if not (starts_with or contains):
        raise BuilderError("DELETE_PARAGRAPH missing anchor",
                           kind="parse", patch_id=patch.id)
    if starts_with:
        matches = _find_paragraphs_starts_with(doc.element.body, starts_with)
        anchor = starts_with
    else:
        matches = _find_paragraphs_containing(doc.element.body, contains)
        anchor = contains
    p = _expect_one(patch.id, "DELETE_PARAGRAPH", "paragraph",
                    anchor, matches)
    _mark_paragraph_as_deleted(p, student_name)
    return ApplyResult(patch_id=patch.id, op="DELETE_PARAGRAPH", status="applied")


def _op_delete_range(doc, patch: Patch, student_name: str) -> ApplyResult:
    start = patch.fields.get("START_PARAGRAPH_STARTS_WITH", "")
    end = patch.fields.get("END_PARAGRAPH_STARTS_WITH", "")
    if not start or not end:
        raise BuilderError("DELETE_RANGE missing START_/END_ anchors",
                           kind="parse", patch_id=patch.id)
    start_matches = _find_paragraphs_starts_with(doc.element.body, start)
    start_p = _expect_one(patch.id, "DELETE_RANGE", "START paragraph",
                          start, start_matches)
    end_matches = _find_paragraphs_starts_with(doc.element.body, end)
    end_p = _expect_one(patch.id, "DELETE_RANGE", "END paragraph",
                        end, end_matches)
    # Verify end is after start
    body = start_p.getparent()
    children = list(body)
    s_idx = children.index(start_p)
    e_idx = children.index(end_p)
    if e_idx < s_idx:
        raise BuilderError(
            "DELETE_RANGE: END appears before START in document.",
            kind="anchor_ambiguous",
            patch_id=patch.id,
        )
    for p in children[s_idx:e_idx + 1]:
        if p.tag == qn("w:p"):
            _mark_paragraph_as_deleted(p, student_name)
    return ApplyResult(patch_id=patch.id, op="DELETE_RANGE", status="applied",
                       detail=f"deleted {e_idx - s_idx + 1} paragraph(s)")


def _op_replace_all(doc, patch: Patch, student_name: str) -> ApplyResult:
    find = patch.fields.get("FIND", "")
    replace = patch.fields.get("REPLACE", "")
    confirm = patch.fields.get("CONFIRM_COUNT")
    if not find:
        raise BuilderError("REPLACE_ALL missing FIND", kind="parse",
                           patch_id=patch.id)

    # Count occurrences across all paragraphs (in document body, optionally
    # restricted by SKIP_IN section list)
    skip_in = patch.fields.get("SKIP_IN") or []
    if isinstance(skip_in, str):
        skip_in = [skip_in]

    find_norm = _norm(find)
    occurrence_count = 0
    paragraphs_with_match = []
    current_section = None
    for p in doc.element.body.iter(qn("w:p")):
        # Track current section by walking headings
        style = _paragraph_style(p)
        ptext = _paragraph_text(p).strip()
        if style and style.startswith("Heading"):
            current_section = ptext
        if current_section in skip_in:
            continue
        p_text = _norm(_paragraph_text(p))
        n = p_text.count(find_norm)
        if n > 0:
            occurrence_count += n
            paragraphs_with_match.append(p)

    if confirm == "unknown":
        # Sonnet-online mode: apply, surface count in audit
        pass
    elif isinstance(confirm, int):
        if confirm != occurrence_count:
            raise BuilderError(
                f"REPLACE_ALL CONFIRM_COUNT mismatch: "
                f"patch says {confirm}, actual count is {occurrence_count}.",
                kind="confirm_count_mismatch",
                patch_id=patch.id,
            )
    else:
        raise BuilderError(
            f"REPLACE_ALL missing CONFIRM_COUNT (integer or 'unknown').",
            kind="parse", patch_id=patch.id,
        )

    # Apply: for each paragraph with a match, do a sub-replace
    for p in paragraphs_with_match:
        _replace_text_in_paragraph(p, find, replace, student_name,
                                    replace_all=True)

    return ApplyResult(
        patch_id=patch.id, op="REPLACE_ALL", status="applied",
        match_count=occurrence_count,
        detail=f"replaced {occurrence_count} occurrence(s)",
    )


def _op_table_cell(doc, patch: Patch, student_name: str) -> ApplyResult:
    # v2.0.0 first cut: find by caption_starts_with, then by ROW_MATCH column+value
    anchor = patch.fields.get("TABLE_ANCHOR", "")
    row_match = patch.fields.get("ROW_MATCH") or {}
    col_target = patch.fields.get("COL", "")
    action = patch.fields.get("ACTION", "REPLACE")
    content = patch.fields.get("CONTENT", "")

    # Find table by adjacent caption
    table_elem = _find_table_by_anchor(doc, anchor)
    if table_elem is None:
        raise BuilderError(
            f"TABLE_CELL anchor not found: {anchor[:80]!r}",
            kind="anchor_missing", patch_id=patch.id,
        )

    # Locate cell via row_match + COL
    row_idx, col_idx = _locate_cell_in_table(table_elem, row_match, col_target)
    if row_idx is None:
        raise BuilderError(
            f"TABLE_CELL ROW_MATCH not found: {row_match!r}",
            kind="anchor_missing", patch_id=patch.id,
        )

    rows = table_elem.findall(qn("w:tr"))
    row = rows[row_idx]
    cells = row.findall(qn("w:tc"))
    cell = cells[col_idx]

    if action == "REPLACE":
        # Wipe cell content, insert new
        for child in list(cell):
            if child.tag in (qn("w:p"),):
                cell.remove(child)
        new_p = OxmlElement("w:p")
        for r in _make_runs_for_text(content):
            new_p.append(r)
        cell.append(new_p)
    elif action == "APPEND":
        new_p = OxmlElement("w:p")
        for r in _make_runs_for_text(content):
            new_p.append(r)
        cell.append(new_p)
    else:
        raise BuilderError(f"TABLE_CELL unknown ACTION: {action!r}",
                           kind="parse", patch_id=patch.id)

    return ApplyResult(patch_id=patch.id, op="TABLE_CELL", status="applied")


def _op_insert_table(doc, patch: Patch, student_name: str) -> ApplyResult:
    """Insert a brand-new table after a heading."""
    heading = patch.fields.get("AFTER_HEADING", "")
    caption = patch.fields.get("CAPTION", "")
    headers = patch.fields.get("HEADERS") or []
    rows = patch.fields.get("ROWS") or []
    if not heading or not headers:
        raise BuilderError("INSERT_TABLE missing AFTER_HEADING or HEADERS",
                           kind="parse", patch_id=patch.id)

    matches = []
    for p in doc.element.body.iter(qn("w:p")):
        if _norm(_paragraph_text(p)) == _norm(heading):
            matches.append(p)
    p = _expect_one(patch.id, "INSERT_TABLE", "heading", heading, matches)

    # Build a w:tbl element
    tbl = _build_table_xml(caption, headers, rows, student_name)
    p.addnext(tbl)
    return ApplyResult(patch_id=patch.id, op="INSERT_TABLE", status="applied")


def _op_figure_regenerated(doc, patch: Patch, student_name: str) -> ApplyResult:
    """No-op on the docx — this is an audit flag only."""
    figure = patch.fields.get("FIGURE", "")
    new_file = patch.fields.get("NEW_FILE", "")
    return ApplyResult(
        patch_id=patch.id, op="FIGURE_REGENERATED", status="applied",
        detail=f"flagged for audit: {figure} → {new_file}",
    )


def _op_ref_list(doc, patch: Patch, student_name: str,
                 renumber_map_accumulator: dict[int, int]) -> ApplyResult:
    """Apply REF_LIST actions.

    Locked v2.0.0 semantics:
      ADD          — parse leading "N." prefix, insert at position N in ref list
      REPLACE_AT N — swap entry at position N in place
      RENUMBER     — explicit MAP {old: new}, applied to body [[CITE:N]] markers
                     (citation renumbering happens AFTER all text patches)
    """
    actions = patch.fields.get("ACTIONS") or []
    if not actions:
        raise BuilderError("REF_LIST missing ACTIONS list",
                           kind="parse", patch_id=patch.id)

    ref_paragraphs = _find_reference_paragraphs(doc)
    detail_parts = []

    for action_block in actions:
        if not isinstance(action_block, dict):
            continue
        if "ADD" in action_block:
            entry = str(action_block["ADD"])
            num = _parse_ref_number_prefix(entry)
            if num is None:
                raise BuilderError(
                    f"REF_LIST ADD entry missing leading 'N.' number prefix: "
                    f"{entry[:60]!r}",
                    kind="parse", patch_id=patch.id,
                )
            _insert_reference_at_position(doc, ref_paragraphs, num, entry,
                                          student_name)
            detail_parts.append(f"ADD {num}")
        elif "REPLACE_AT" in action_block:
            pos = int(action_block["REPLACE_AT"])
            new_text = str(action_block.get("NEW", ""))
            if not new_text:
                raise BuilderError(
                    f"REF_LIST REPLACE_AT {pos} missing NEW",
                    kind="parse", patch_id=patch.id,
                )
            _replace_reference_at_position(doc, ref_paragraphs, pos, new_text,
                                            student_name)
            detail_parts.append(f"REPLACE_AT {pos}")
        elif "RENUMBER" in action_block:
            map_val = action_block.get("MAP") or action_block.get("RENUMBER")
            if not isinstance(map_val, dict):
                raise BuilderError(
                    "REF_LIST RENUMBER requires MAP: {old: new} dict",
                    kind="parse", patch_id=patch.id,
                )
            # Accumulate the map — applied to body citations after all patches
            for old_k, new_v in map_val.items():
                renumber_map_accumulator[int(old_k)] = int(new_v)
            detail_parts.append(f"RENUMBER {map_val}")
        else:
            raise BuilderError(
                f"REF_LIST unknown action: {action_block!r}",
                kind="parse", patch_id=patch.id,
            )

    return ApplyResult(
        patch_id=patch.id, op="REF_LIST", status="applied",
        detail="; ".join(detail_parts),
    )


# ─── Paragraph-content manipulation ────────────────────────────────────────

def _replace_text_in_paragraph(p_elem, find: str, replace: str,
                                student_name: str, *,
                                replace_all: bool = False) -> None:
    """Replace `find` text inside paragraph with `replace` (with [[CITE:N]]
    markers resolved). The original `find` content is wrapped in <w:del>;
    the new content is wrapped in <w:ins>.

    v2.0.0 implementation: paragraph-level swap. Every text-bearing child
    of the paragraph (direct <w:r>, runs inside <w:ins> from earlier patches,
    runs inside <w:hyperlink>) becomes part of one <w:del>, followed by a
    single <w:ins> with the REPLACE text. This produces correct text on
    Accept All Changes but the tracked changes appear at paragraph
    granularity rather than at the surgical FIND span. A future v2.x can
    add a surgical mode that splits runs around the FIND boundaries.
    """
    # Collect all text-bearing children (anything whose subtree contains <w:t>).
    # Excludes <w:pPr> and structural elements without text (bookmarks etc.).
    text_bearing = []
    for child in list(p_elem):
        if child.tag == qn("w:pPr"):
            continue
        if any(t.text for t in child.iter(qn("w:t"))):
            text_bearing.append(child)

    if not text_bearing:
        return  # nothing to replace

    # Deep-copy each text-bearing child for the <w:del> wrapper, then
    # extract their inner <w:r> elements (flattening <w:ins>, <w:hyperlink>).
    runs_for_del = []
    for child in text_bearing:
        child_copy = copy.deepcopy(child)
        if child_copy.tag == qn("w:r"):
            runs_for_del.append(child_copy)
        else:
            # <w:ins> from an earlier patch, <w:hyperlink>, or similar
            # container — pull out its <w:r> children
            for r in child_copy.iter(qn("w:r")):
                runs_for_del.append(r)

    # Remove originals from the paragraph
    for child in text_bearing:
        p_elem.remove(child)

    if not runs_for_del:
        return

    del_block = _wrap_del(runs_for_del, student_name)
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(del_block)
    else:
        p_elem.insert(0, del_block)

    # Append the REPLACE content as <w:ins>
    if replace:
        new_runs = _make_runs_for_text(replace)
        ins_block = _wrap_ins(new_runs, student_name)
        del_block.addnext(ins_block)


def _insert_paragraph_after(p_elem, text: str, student_name: str,
                             style: str | None = None) -> None:
    """Insert a new paragraph (wrapped in <w:ins>) after p_elem."""
    new_p = OxmlElement("w:p")
    if style:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style)
        pPr.append(pStyle)
        new_p.append(pPr)
    # The whole paragraph is an insertion: wrap pPr + runs in <w:ins> by
    # using the inserted-paragraph-mark trick: mark the trailing <w:p>'s
    # paragraph mark via pPr/rPr/<w:ins>, AND wrap runs in <w:ins>.
    # For simplicity: place all runs inside a single <w:ins> block.
    ins = _wrap_ins(_make_runs_for_text(text), student_name)
    new_p.append(ins)

    # Mark the paragraph mark itself as inserted (so Accept-All in Word
    # treats the whole paragraph as inserted, not just its content)
    pPr2 = new_p.find(qn("w:pPr")) or OxmlElement("w:pPr")
    rPr = OxmlElement("w:rPr")
    ins_mark = OxmlElement("w:ins")
    ins_mark.set(qn("w:id"), _next_change_id())
    ins_mark.set(qn("w:author"), student_name)
    ins_mark.set(qn("w:date"),
                  datetime.datetime.now().replace(microsecond=0).isoformat() + "Z")
    rPr.append(ins_mark)
    pPr2.append(rPr)
    if new_p.find(qn("w:pPr")) is None:
        new_p.insert(0, pPr2)

    p_elem.addnext(new_p)


def _insert_heading_after(p_elem, text: str, level: int, student_name: str) -> None:
    """Insert a Heading {level} paragraph after p_elem."""
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), f"Heading{level}")
    pPr.append(pStyle)
    new_p.append(pPr)
    ins = _wrap_ins(_make_runs_for_text(text), student_name)
    new_p.append(ins)
    p_elem.addnext(new_p)


def _mark_paragraph_as_deleted(p_elem, student_name: str) -> None:
    """Mark every visible run in p_elem as deleted. Also mark the paragraph
    mark as deleted so Accept All causes the paragraph to disappear entirely."""
    runs_to_delete = []
    for child in list(p_elem):
        if child.tag == qn("w:r"):
            runs_to_delete.append(child)
        elif child.tag == qn("w:ins"):
            # If a previous patch inserted this content, it's now being
            # un-inserted: just remove the ins block entirely.
            p_elem.remove(child)
    if runs_to_delete:
        del_runs = [copy.deepcopy(r) for r in runs_to_delete]
        for r in runs_to_delete:
            p_elem.remove(r)
        del_block = _wrap_del(del_runs, student_name)
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is not None:
            pPr.addnext(del_block)
        else:
            p_elem.insert(0, del_block)

    # Mark the paragraph mark itself as deleted
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    del_mark = OxmlElement("w:del")
    del_mark.set(qn("w:id"), _next_change_id())
    del_mark.set(qn("w:author"), student_name)
    del_mark.set(qn("w:date"),
                  datetime.datetime.now().replace(microsecond=0).isoformat() + "Z")
    rPr.append(del_mark)


# ─── Reference list manipulation ───────────────────────────────────────────

def _find_reference_paragraphs(doc) -> list:
    """Return the paragraph elements that comprise the reference list,
    in order. Detection: a paragraph whose style is Heading* with text
    matching 'References' (case-insensitive), followed by paragraphs
    starting with `N.`"""
    body = doc.element.body
    in_refs = False
    refs = []
    for p in body.iter(qn("w:p")):
        text = _paragraph_text(p).strip()
        style = _paragraph_style(p) or ""
        if not in_refs:
            if style.startswith("Heading") and _norm(text).lower().startswith("references"):
                in_refs = True
                continue
            # Also accept plain "References" paragraph (no heading style)
            if _norm(text).lower() == "references":
                in_refs = True
                continue
            continue
        if not text:
            continue
        if _REF_NUMBER_PREFIX_RE.match(text):
            refs.append(p)
        else:
            # Non-numbered line in refs section — likely end of list
            break
    return refs


def _parse_ref_number_prefix(text: str) -> int | None:
    m = _REF_NUMBER_PREFIX_RE.match(text)
    if m:
        return int(m.group(1))
    return None


def _insert_reference_at_position(doc, ref_paragraphs: list, num: int,
                                   entry_text: str, student_name: str) -> None:
    """Insert a new reference entry at position `num`. If num is past the
    end of the existing list, append after the last."""
    # ref_paragraphs is the current list (possibly mutated by previous adds).
    # Re-locate to be safe.
    refs = _find_reference_paragraphs(doc)
    new_p = OxmlElement("w:p")
    ins = _wrap_ins(_make_runs_for_text(entry_text), student_name)
    new_p.append(ins)
    # Mark paragraph mark as inserted too
    pPr = OxmlElement("w:pPr")
    rPr = OxmlElement("w:rPr")
    ins_mark = OxmlElement("w:ins")
    ins_mark.set(qn("w:id"), _next_change_id())
    ins_mark.set(qn("w:author"), student_name)
    ins_mark.set(qn("w:date"),
                  datetime.datetime.now().replace(microsecond=0).isoformat() + "Z")
    rPr.append(ins_mark)
    pPr.append(rPr)
    new_p.insert(0, pPr)

    if num <= len(refs):
        # Insert before refs[num - 1]
        refs[num - 1].addprevious(new_p)
    elif refs:
        refs[-1].addnext(new_p)
    else:
        # No existing references — append after the References heading
        body = doc.element.body
        heading = None
        for p in body.iter(qn("w:p")):
            if _norm(_paragraph_text(p)).lower() == "references":
                heading = p
                break
        if heading is not None:
            heading.addnext(new_p)
        else:
            body.append(new_p)


def _replace_reference_at_position(doc, ref_paragraphs: list, pos: int,
                                    new_text: str, student_name: str) -> None:
    refs = _find_reference_paragraphs(doc)
    if pos < 1 or pos > len(refs):
        raise BuilderError(
            f"REF_LIST REPLACE_AT {pos}: only {len(refs)} references in list.",
            kind="anchor_missing",
        )
    target = refs[pos - 1]
    old_text = _paragraph_text(target)
    _replace_text_in_paragraph(target, old_text, new_text, student_name)


# ─── Table helpers ─────────────────────────────────────────────────────────

def _find_table_by_anchor(doc, anchor: str) -> Any:
    """Find <w:tbl> whose preceding caption matches `anchor` (caption_starts_with)."""
    if anchor.startswith("caption_starts_with="):
        prefix = anchor.split("=", 1)[1].strip('"').strip("'")
    else:
        prefix = anchor
    body = doc.element.body
    prev_text = ""
    for child in body.iter():
        if child.tag == qn("w:p"):
            prev_text = _paragraph_text(child).strip()
        elif child.tag == qn("w:tbl") and prev_text and _norm(prev_text).startswith(_norm(prefix)):
            return child
    return None


def _locate_cell_in_table(table_elem, row_match: dict, col_target: str) -> tuple:
    rows = table_elem.findall(qn("w:tr"))
    if not rows:
        return None, None
    # First row = header. Build column index map.
    header_cells = rows[0].findall(qn("w:tc"))
    col_to_idx = {}
    for i, cell in enumerate(header_cells):
        text = "".join(t.text or "" for t in cell.iter(qn("w:t"))).strip()
        col_to_idx[_norm(text).lower()] = i

    match_col = row_match.get("column", "").lower()
    match_val = row_match.get("value", "")
    match_col_idx = col_to_idx.get(_norm(match_col).lower())
    target_col_idx = col_to_idx.get(_norm(col_target).lower())
    if match_col_idx is None or target_col_idx is None:
        return None, None

    for ri, row in enumerate(rows[1:], start=1):
        cells = row.findall(qn("w:tc"))
        if match_col_idx >= len(cells):
            continue
        cell_text = "".join(t.text or "" for t in cells[match_col_idx].iter(qn("w:t"))).strip()
        if _norm(cell_text).lower() == _norm(match_val).lower():
            return ri, target_col_idx
    return None, None


def _build_table_xml(caption: str, headers: list, rows: list,
                      student_name: str) -> OxmlElement:
    """Build a minimal <w:tbl> with headers + rows. Caption is emitted as a
    separate <w:p> before the table (caller inserts as appropriate)."""
    tbl = OxmlElement("w:tbl")
    # tblPr (minimal)
    tblPr = OxmlElement("w:tblPr")
    tbl.append(tblPr)
    # Build header row
    header_tr = OxmlElement("w:tr")
    for h in headers:
        tc = OxmlElement("w:tc")
        p = OxmlElement("w:p")
        for r in _make_runs_for_text(str(h)):
            p.append(r)
        tc.append(p)
        header_tr.append(tc)
    tbl.append(header_tr)
    # Data rows
    for row in rows:
        tr = OxmlElement("w:tr")
        for cell in row:
            tc = OxmlElement("w:tc")
            p = OxmlElement("w:p")
            for r in _make_runs_for_text(str(cell)):
                p.append(r)
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    return tbl


# ─── Renumber map application ──────────────────────────────────────────────

def _apply_renumber_map(doc, renumber_map: dict[int, int]) -> None:
    """Walk all superscript citation runs in the docx and remap N → renumber_map[N].
    Also remap [[CITE:N]] markers that haven't been resolved yet (defensive)."""
    if not renumber_map:
        return
    body = doc.element.body
    for r in body.iter(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            continue
        va = rPr.find(qn("w:vertAlign"))
        if va is None or va.get(qn("w:val")) != "superscript":
            continue
        t = r.find(qn("w:t"))
        if t is None or not t.text:
            continue
        if not re.fullmatch(r"[\d,\s]+", t.text):
            continue
        # Remap
        new_nums = []
        for piece in re.split(r",", t.text):
            piece = piece.strip()
            if piece.isdigit():
                old = int(piece)
                new_nums.append(str(renumber_map.get(old, old)))
        t.text = ",".join(new_nums)


# ─── Standard-Clean derivation ─────────────────────────────────────────────

def _derive_standard_clean(tracked_path: Path, clean_path: Path) -> None:
    """Read tracked docx, strip <w:del>, unwrap <w:ins>, write clean."""
    doc = Document(str(tracked_path))
    body = doc.element.body

    # Pass 1: remove paragraphs whose paragraph-mark is marked <w:del>
    to_remove = []
    for p in body.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            rPr = pPr.find(qn("w:rPr"))
            if rPr is not None and rPr.find(qn("w:del")) is not None:
                to_remove.append(p)
    for p in to_remove:
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    # Pass 2: remove all <w:del> elements (their content was deleted)
    for d in list(body.iter(qn("w:del"))):
        parent = d.getparent()
        if parent is not None:
            parent.remove(d)

    # Pass 3: unwrap all <w:ins> elements (their content becomes plain)
    for ins in list(body.iter(qn("w:ins"))):
        parent = ins.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ins)
        for child in list(ins):
            ins.remove(child)
            parent.insert(idx, child)
            idx += 1
        parent.remove(ins)

    # Pass 4: strip residual <w:ins> paragraph-mark markers in pPr/rPr
    for rPr in body.iter(qn("w:rPr")):
        for ins_mark in list(rPr.findall(qn("w:ins"))):
            rPr.remove(ins_mark)

    doc.save(str(clean_path))


# ─── Build orchestrator ────────────────────────────────────────────────────

# Op routing table — populated below after function definitions
_TEXT_OPS = ("FIND_REPLACE", "APPEND_TO_PARAGRAPH", "INSERT_AFTER_PARAGRAPH",
             "INSERT_AFTER_HEADING", "REPLACE_HEADING",
             "DELETE_PARAGRAPH", "DELETE_RANGE")
_TABLE_OPS = ("TABLE_CELL", "INSERT_TABLE")
_SCOPE_OPS = ("REPLACE_ALL",)
_REF_OPS = ("REF_LIST",)
_FLAG_OPS = ("FIGURE_REGENERATED",)

_OP_DISPATCH: dict[str, Callable] = {
    "FIND_REPLACE": _op_find_replace,
    "APPEND_TO_PARAGRAPH": _op_append_to_paragraph,
    "INSERT_AFTER_PARAGRAPH": _op_insert_after_paragraph,
    "INSERT_AFTER_HEADING": _op_insert_after_heading,
    "REPLACE_HEADING": _op_replace_heading,
    "DELETE_PARAGRAPH": _op_delete_paragraph,
    "DELETE_RANGE": _op_delete_range,
    "REPLACE_ALL": _op_replace_all,
    "TABLE_CELL": _op_table_cell,
    "INSERT_TABLE": _op_insert_table,
    "FIGURE_REGENERATED": _op_figure_regenerated,
}


def apply_patches(revision: RevisionDoc, original_docx_path: str,
                  output_tracked_path: str) -> ApplyReport:
    """Apply revision.patches to original_docx, save as output_tracked_path.

    Returns ApplyReport with per-patch results + warnings.
    Order per v2.2 Phase 4: text → table → scope → REF_LIST.
    """
    doc = Document(str(original_docx_path))
    student = revision.student_name

    report = ApplyReport()
    renumber_map: dict[int, int] = {}

    # Bucket patches by category, preserving document order within each
    text_patches = [p for p in revision.patches if p.op in _TEXT_OPS]
    table_patches = [p for p in revision.patches if p.op in _TABLE_OPS]
    scope_patches = [p for p in revision.patches if p.op in _SCOPE_OPS]
    ref_patches = [p for p in revision.patches if p.op in _REF_OPS]
    flag_patches = [p for p in revision.patches if p.op in _FLAG_OPS]
    unknown_patches = [p for p in revision.patches
                        if p.op not in _OP_DISPATCH]

    for unknown in unknown_patches:
        report.results.append(ApplyResult(
            patch_id=unknown.id, op=unknown.op, status="failed",
            detail=f"unknown op {unknown.op!r} — not in v2.2 vocabulary",
        ))

    def run(patches_list, name):
        for patch in patches_list:
            handler = _OP_DISPATCH[patch.op]
            try:
                if patch.op == "REF_LIST":
                    res = _op_ref_list(doc, patch, student, renumber_map)
                else:
                    res = handler(doc, patch, student)
                report.results.append(res)
            except BuilderError as exc:
                report.results.append(ApplyResult(
                    patch_id=patch.id, op=patch.op,
                    status="failed", detail=str(exc),
                ))
                # Strict mode: stop on first failure so the mentor sees
                # exactly which patch broke
                raise

    # Order per Phase 4
    run(text_patches, "text")
    run(table_patches, "table")
    run(scope_patches, "scope")
    run(ref_patches, "ref")
    run(flag_patches, "flag")

    # Apply renumber map (post all text patches, per spec)
    if renumber_map:
        _apply_renumber_map(doc, renumber_map)
        report.warnings.append(
            f"Applied citation renumber map: {renumber_map}"
        )

    # Set author attribution
    doc.core_properties.author = student
    doc.core_properties.last_modified_by = student
    if revision.manuscript_title:
        doc.core_properties.title = revision.manuscript_title

    doc.save(str(output_tracked_path))
    return report


# ─── Filename helper ───────────────────────────────────────────────────────

_FILENAME_NONALNUM = re.compile(r"[^a-z0-9]+")
_FILENAME_TRIM = re.compile(r"^-+|-+$")


def sanitize_filename_stem(title: str, *, cap: int = 80) -> str:
    if not title:
        return "untitled"
    stem = title.lower()
    stem = _FILENAME_NONALNUM.sub("-", stem)
    stem = _FILENAME_TRIM.sub("", stem)
    if len(stem) > cap:
        stem = stem[:cap].rstrip("-")
    return stem or "untitled"


# ─── Response letter + Student handoff (carried over from v1) ──────────────

_ACTION_HEADER_RE = re.compile(r"^###\s+Action\s+([A-Za-z0-9]+)\s*(?:—|-)\s*(.+?)\s*$")
_COMMENT_HEADER_RE = re.compile(r"^###\s+Comment\s+(R\d+\.\d+)\s*(?:—|-)\s*(.+?)\s*$")
_KV_RE = re.compile(r"^([A-Za-z][A-Za-z0-9 _-]*?)\s*:\s*(.*)$")
_REVIEWER_QUOTE_RE = re.compile(
    r'^Reviewer\s+(R\d+\.\d+)\s*:\s*"?(.*?)"?\s*$', re.IGNORECASE
)


def _parse_revision_log(path: str | Path | None) -> dict:
    if path is None:
        return {"actions": [], "pushed_back": []}
    p = Path(path)
    if not p.exists():
        return {"actions": [], "pushed_back": []}
    text = p.read_text(encoding="utf-8")
    lines = text.splitlines()

    actions: list[dict] = []
    pushed: list[dict] = []
    current: dict | None = None
    current_kind: str | None = None

    def _commit():
        nonlocal current, current_kind
        if current is None:
            return
        if current_kind == "action":
            actions.append(current)
        elif current_kind == "comment":
            pushed.append(current)
        current = None
        current_kind = None

    for ln in lines:
        if ln.startswith("## ") and not ln.startswith("### "):
            _commit()
            continue
        m_action = _ACTION_HEADER_RE.match(ln)
        if m_action:
            _commit()
            current = {"id": m_action.group(1), "summary": m_action.group(2).strip(),
                       "comments": [], "reviewer_text": {}}
            current_kind = "action"
            continue
        m_comment = _COMMENT_HEADER_RE.match(ln)
        if m_comment:
            _commit()
            current = {"id": m_comment.group(1), "summary": m_comment.group(2).strip(),
                       "reviewer_text": {}}
            current_kind = "comment"
            continue
        if current is None:
            continue
        m_quote = _REVIEWER_QUOTE_RE.match(ln)
        if m_quote:
            current["reviewer_text"][m_quote.group(1)] = m_quote.group(2).strip()
            continue
        m_kv = _KV_RE.match(ln)
        if m_kv:
            key = m_kv.group(1).strip().lower().replace(" ", "_")
            val = m_kv.group(2).strip()
            if key == "comments":
                current["comments"] = [c.strip() for c in val.split(",") if c.strip()]
            else:
                current[key] = val
    _commit()
    return {"actions": actions, "pushed_back": pushed}


_QUOTE_PLACEHOLDER = "[REVIEWER COMMENT QUOTE — INSERT FROM PDF]"


def _reviewer_index_from_id(comment_id: str) -> int | None:
    m = re.match(r"^R(\d+)\.\d+$", comment_id)
    return int(m.group(1)) if m else None


def _comment_sort_key(comment: dict) -> tuple:
    m = re.match(r"^R(\d+)\.(\d+)$", comment.get("comment_id", ""))
    return (int(m.group(1)), int(m.group(2))) if m else (999, 999)


def build_response_letter(revision: RevisionDoc,
                          revision_log_md_path: str | None,
                          output_path: str) -> None:
    log = _parse_revision_log(revision_log_md_path)
    actions = log["actions"]
    pushed = log["pushed_back"]

    comments_by_reviewer: dict[int, list[dict]] = {}
    for action in actions:
        for cid in action.get("comments", []):
            idx = _reviewer_index_from_id(cid)
            if idx is None:
                continue
            comments_by_reviewer.setdefault(idx, []).append({
                "comment_id": cid,
                "verbatim": action["reviewer_text"].get(cid, ""),
                "response": _craft_response_line(
                    action.get("classification", ""),
                    action.get("change_summary", ""),
                ),
                "location": action.get("location") or
                            "No manuscript change; see response above.",
                "kind": "action",
            })
    for pb in pushed:
        cid = pb.get("id", "")
        idx = _reviewer_index_from_id(cid)
        if idx is None:
            continue
        comments_by_reviewer.setdefault(idx, []).append({
            "comment_id": cid,
            "verbatim": pb["reviewer_text"].get(cid, ""),
            "response": pb.get("push_back_defense", ""),
            "location": "No manuscript change; see response above.",
            "kind": "pushback",
        })
    for idx in comments_by_reviewer:
        comments_by_reviewer[idx].sort(key=_comment_sort_key)

    doc = Document()
    doc.core_properties.author = revision.student_name
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)
    doc.add_heading("Reviewer Response Letter", level=0)
    doc.add_paragraph("Dear Editor and Reviewers,")
    doc.add_paragraph(
        "We thank the reviewers for their constructive and detailed feedback. "
        "Below we address each comment point by point. Changes to the "
        "manuscript are marked in red text. Reviewer comments are quoted "
        "verbatim."
    )

    if not comments_by_reviewer:
        warn = doc.add_paragraph()
        wr = warn.add_run(
            "[PLACEHOLDER LETTER — no revision-log.md provided. Fill in "
            "comments, responses, and change locations manually.]"
        )
        wr.font.color.rgb = RED
        wr.italic = True
    else:
        for idx in sorted(comments_by_reviewer):
            doc.add_heading(f"Reviewer {idx}", level=1)
            for c in comments_by_reviewer[idx]:
                doc.add_heading(f"Comment {c['comment_id']}", level=2)
                q = doc.add_paragraph()
                qr = q.add_run(f'"{c["verbatim"]}"' if c["verbatim"]
                                else _QUOTE_PLACEHOLDER)
                qr.italic = True
                rp = doc.add_paragraph()
                rp.add_run("Response: ").bold = True
                rp.add_run(c["response"] or "[FILL IN RESPONSE]")
                lp = doc.add_paragraph()
                lp.add_run("Change location: ").bold = True
                lp.add_run(c["location"])

    doc.add_heading("Summary of major changes", level=1)
    if actions:
        for a in actions:
            sm = a.get("change_summary", "")
            s = a.get("summary", "")
            doc.add_paragraph(f"{s}. {sm}" if sm else s or "")
    else:
        doc.add_paragraph("[FILL IN MAJOR CHANGES SUMMARY]")

    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(revision.student_name)
    doc.save(output_path)


def _craft_response_line(classification: str, change_summary: str) -> str:
    if not change_summary:
        return "[FILL IN — change summary was missing from the log]"
    cls = (classification or "").lower()
    if "push" in cls:
        return change_summary
    if "out of scope" in cls:
        return f"We respectfully note this. {change_summary}"
    return f"We have made this change. {change_summary}"


def build_student_handoff(revision: RevisionDoc,
                          revision_log_md_path: str | None,
                          output_path: str) -> None:
    log = _parse_revision_log(revision_log_md_path)
    actions = log["actions"]
    pushed = log["pushed_back"]
    stem = sanitize_filename_stem(revision.manuscript_title)

    lines: list[str] = []
    lines.append(f"# {revision.student_name} — NHSJS Revision Handoff")
    lines.append("")
    lines.append(f"_Manuscript: {revision.manuscript_title}_")
    lines.append("")
    lines.append("## 1. What the reviewer said")
    lines.append("")
    by_section: dict[str, list[str]] = {}
    for a in actions:
        sec = a.get("section", "").strip() or "General"
        by_section.setdefault(sec, []).append(a.get("summary", ""))
    if by_section:
        for sec, summaries in by_section.items():
            if len(summaries) == 1:
                lines.append(f"- {sec}: {summaries[0]}")
            else:
                lines.append(f"- {sec}: {len(summaries)} comments addressed")
        if pushed:
            lines.append(f"- {len(pushed)} comment(s) respectfully declined")
    else:
        lines.append("- [FILL IN — 3–5 themes from reviewer decision]")
    lines.append("")
    lines.append("## 2. What we changed and how")
    lines.append("")
    if actions:
        for a in actions:
            sm = a.get("change_summary", "").strip()
            loc = a.get("location", "").strip()
            if sm:
                lines.append(f"- {sm}" + (f" (See: {loc}.)" if loc else ""))
    else:
        lines.append("- [FILL IN — major changes with location]")
    if pushed:
        lines.append("")
        lines.append("Pushed back on:")
        for pb in pushed:
            lines.append(f"- {pb.get('summary', pb.get('id', '?'))} (see Response-Letter.docx)")
    lines.append("")
    lines.append("## 3. What you need to do")
    lines.append("")
    lines.append("- Open `Standard-Tracked.docx`, read the red text top to "
                 "bottom, and click Accept or Reject as you go. Most should "
                 "be Accept — the clean version assumes everything is accepted.")
    lines.append("")
    lines.append("## 4. What's needed to send")
    lines.append("")
    lines.append("- [ ] Read through `Standard-Clean.docx` and confirm it says what you want")
    lines.append("- [ ] Confirm your name is on the docx as the author (it should be)")
    lines.append(f"- [ ] Reply to the NHSJS email with these three files attached: "
                 f"`{stem}-Standard-Tracked.docx`, `{stem}-Online.docx`, "
                 f"`Response-Letter.docx`")
    lines.append("- [ ] Confirm submission receipt within 24 hours")
    lines.append("")
    lines.append("Any questions, ping me.")
    lines.append("")
    Path(output_path).write_text("\n".join(lines), encoding="utf-8")


# ─── FILES.md ──────────────────────────────────────────────────────────────

_FILES_MD_TEMPLATE = """\
# Submission Files — {title}

Generated by nhsjs-tools v{version} on {date}.

| File | What it is |
|------|-----------|
| `{tracked_name}` | Standard-format manuscript with tracked changes. Insertions in red bold (real OOXML `<w:ins>`), deletions in red strikethrough (real OOXML `<w:del>`). Open in Word to Accept/Reject. Send to NHSJS so reviewers can see exactly what changed. |
| `{clean_name}` | Standard-format manuscript with all changes accepted. The version of record for the resubmission. |
| `{online_name}` | Online-publication format. Citations are full-text in `((double parens))` instead of numeric superscript. Send to NHSJS alongside the Standard version. |
| `Response-Letter.docx` | Point-by-point reviewer response letter. |
| `Student-Handoff.md` | Internal handoff for the student. Not for submission. |

Author attribution: **{student_name}**.
"""


# ─── build_all ─────────────────────────────────────────────────────────────

def build_all(revision_md_path: str,
              original_docx_path: str,
              revision_log_md_path: str | None,
              output_dir: str) -> dict[str, str]:
    """Run the full v2.0.0 build pipeline.

    Args:
        revision_md_path: path to revision.md (v2.2 patch syntax)
        original_docx_path: path to the original Standard docx
        revision_log_md_path: optional revision-log.md for letter + handoff
        output_dir: where to write the 5 outputs + FILES.md

    Returns:
        {kind: path} dict mapping output kind to absolute path.

    Raises:
        BuilderError on strict-mode failures.
    """
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)

    revision = parse_revision(revision_md_path)
    stem = sanitize_filename_stem(revision.manuscript_title)

    tracked_path = out / f"{stem}-Standard-Tracked.docx"
    clean_path = out / f"{stem}-Standard-Clean.docx"
    online_path = out / f"{stem}-Online.docx"
    letter_path = out / "Response-Letter.docx"
    handoff_path = out / "Student-Handoff.md"
    files_md_path = out / "FILES.md"

    # Apply patches → Standard-Tracked
    report = apply_patches(revision, original_docx_path, str(tracked_path))

    # Derive Standard-Clean
    _derive_standard_clean(tracked_path, clean_path)

    # Pipe Standard-Clean → Standard→Online
    from nhsjs_standard_to_online import convert_standard_to_online
    convert_standard_to_online(str(clean_path), str(online_path))

    # Letter + handoff
    build_response_letter(revision, revision_log_md_path, str(letter_path))
    build_student_handoff(revision, revision_log_md_path, str(handoff_path))

    # FILES.md
    files_md_path.write_text(
        _FILES_MD_TEMPLATE.format(
            title=revision.manuscript_title,
            version=_VERSION,
            date=datetime.date.today().isoformat(),
            student_name=revision.student_name,
            tracked_name=tracked_path.name,
            clean_name=clean_path.name,
            online_name=online_path.name,
        ),
        encoding="utf-8",
    )

    result = {
        "standard_tracked": str(tracked_path.resolve()),
        "standard_clean":   str(clean_path.resolve()),
        "online":           str(online_path.resolve()),
        "response_letter":  str(letter_path.resolve()),
        "student_handoff":  str(handoff_path.resolve()),
        "files_md":         str(files_md_path.resolve()),
    }
    # Attach the apply report so callers (Streamlit, audit) can surface it
    result["_apply_report"] = report  # type: ignore[assignment]
    return result


# ─── CLI ───────────────────────────────────────────────────────────────────

def _cli():
    parser = argparse.ArgumentParser(
        description="Build NHSJS revision outputs from revision.md + original docx (v2.0.0).",
    )
    parser.add_argument("revision_md", help="Path to revision.md")
    parser.add_argument("original_docx", help="Path to the original Standard .docx")
    parser.add_argument("output_dir", nargs="?", default="submit",
                        help="Output directory (default: submit/)")
    parser.add_argument("--log", metavar="FILE",
                        help="Optional revision-log.md for letter + handoff")
    args = parser.parse_args()

    try:
        result = build_all(args.revision_md, args.original_docx,
                           args.log, args.output_dir)
    except BuilderError as exc:
        print(f"error ({exc.kind}): {exc}", file=sys.stderr)
        if exc.patch_id:
            print(f"  in patch: {exc.patch_id}", file=sys.stderr)
        if exc.line:
            print(f"  at line: {exc.line}", file=sys.stderr)
        raise SystemExit(2)

    for kind, path in result.items():
        if kind.startswith("_"):
            continue
        print(f"  {kind:<20}  {path}")


if __name__ == "__main__":
    _cli()
