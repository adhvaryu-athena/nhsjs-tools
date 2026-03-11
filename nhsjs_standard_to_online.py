"""
NHSJS → NHSJS Online Citation Converter
========================================
Converts a standard NHSJS Word document (superscript numbered citations)
to the NHSJS Online publication format (full citations in double parentheses).

Usage:
    python nhsjs_converter.py input.docx output.docx
    python nhsjs_converter.py input.docx output.docx --refs refs.txt

Requirements:
    pip install python-docx lxml
"""

import sys
import re
import copy
import argparse
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── XML helpers ──────────────────────────────────────────────────────────────

def is_superscript(r_elem):
    rPr = r_elem.find(qn('w:rPr'))
    if rPr is not None:
        va = rPr.find(qn('w:vertAlign'))
        if va is not None and va.get(qn('w:val')) == 'superscript':
            return True
    return False

def get_run_text(r_elem):
    t = r_elem.find(qn('w:t'))
    return t.text if (t is not None and t.text) else ''

def make_plain_run(text, base_rPr=None):
    r = OxmlElement('w:r')
    if base_rPr is not None:
        new_rPr = copy.deepcopy(base_rPr)
        for va in new_rPr.findall(qn('w:vertAlign')):
            new_rPr.remove(va)
        for color in new_rPr.findall(qn('w:color')):
            new_rPr.remove(color)
        r.append(new_rPr)
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    return r

def make_superscript_comma():
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    va = OxmlElement('w:vertAlign')
    va.set(qn('w:val'), 'superscript')
    rPr.append(va)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = ','
    r.append(t)
    return r


# ─── Reference parsers ────────────────────────────────────────────────────────

def _parse_numbered_lines(lines):
    """Parse '1. Citation…' lines into {int: str}. Handles multi-line entries."""
    citation_map = {}
    current_num = None
    current_body = []

    for line in lines:
        line = line.strip()
        if not line:
            continue
        m = re.match(r'^(\d+)[.)]\s+(.+)', line)
        if m:
            if current_num is not None:
                body = ' '.join(current_body).strip()
                if not body.endswith('.'):
                    body += '.'
                citation_map[current_num] = body
            current_num = int(m.group(1))
            current_body = [m.group(2).strip()]
        elif current_num is not None:
            current_body.append(line)  # wrapped continuation line

    if current_num is not None:
        body = ' '.join(current_body).strip()
        if not body.endswith('.'):
            body += '.'
        citation_map[current_num] = body

    return citation_map


def parse_references_from_doc(doc):
    """
    Find the References section in the docx and parse numbered entries.
    Tries: (1) heading style containing 'reference', (2) plain paragraph 'References'.
    Returns (citation_map, method_str, ref_start_idx).
    """
    ref_start_idx = None
    method = None

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name.lower() if p.style else ''
        if 'heading' in style and 'reference' in text.lower():
            ref_start_idx = i
            method = f'heading style "{p.style.name}"'
            break
        if re.match(r'^references\s*:?\s*$', text, re.IGNORECASE):
            ref_start_idx = i
            method = 'plain paragraph "References"'
            break

    if ref_start_idx is None:
        return {}, None, None

    lines = [doc.paragraphs[i].text for i in range(ref_start_idx + 1, len(doc.paragraphs))]
    return _parse_numbered_lines(lines), method, ref_start_idx


def parse_references_from_txt(path):
    """Parse a plain .txt file of numbered references."""
    with open(path, encoding='utf-8', errors='replace') as f:
        lines = f.readlines()
    # Skip optional 'References' header line
    start = 0
    for i, line in enumerate(lines):
        if re.match(r'^references\s*:?\s*$', line.strip(), re.IGNORECASE):
            start = i + 1
            break
    return _parse_numbered_lines(lines[start:])


# ─── Paragraph processor ──────────────────────────────────────────────────────

def process_paragraph(p_elem, citation_map):
    children = list(p_elem)
    groups = []
    i = 0
    while i < len(children):
        child = children[i]
        if child.tag != qn('w:r') or not is_superscript(child):
            i += 1
            continue
        run_text = get_run_text(child)
        digits = re.findall(r'\d+', run_text)
        if not digits:
            i += 1
            continue
        nums = [int(d) for d in digits]
        j = i + 1
        while j < len(children):
            nc = children[j]
            if nc.tag != qn('w:r') or not is_superscript(nc):
                break
            nt = get_run_text(nc)
            more_digits = re.findall(r'\d+', nt)
            if more_digits or re.fullmatch(r'[\s,]+', nt):
                nums.extend(int(d) for d in more_digits)
                j += 1
            else:
                break
        groups.append((i, j, nums))
        i = j

    if not groups:
        return []

    missing = []
    for start, end, nums in reversed(groups):
        base_rPr = children[start].find(qn('w:rPr'))
        replacement = []
        for k, num in enumerate(nums):
            if k > 0:
                replacement.append(make_superscript_comma())
            citation_text = citation_map.get(num)
            if citation_text:
                replacement.append(make_plain_run(f" (({citation_text}))", base_rPr))
            else:
                replacement.append(make_plain_run(f"[Citation {num} not found]", base_rPr))
                missing.append(num)
        insert_pos = p_elem.index(children[start])
        for r in children[start:end]:
            p_elem.remove(r)
        for offset, node in enumerate(replacement):
            p_elem.insert(insert_pos + offset, node)

    return missing


# ─── Main ─────────────────────────────────────────────────────────────────────

def convert(input_path, output_path, refs_txt=None):
    print(f"\n📄  Loading: {input_path}")
    doc = Document(input_path)

    # Resolve citation map
    if refs_txt:
        citation_map = parse_references_from_txt(refs_txt)
        ref_start_idx = None
        print(f"📋  Using references from: {refs_txt}")
        method = "external .txt file"
    else:
        citation_map, method, ref_start_idx = parse_references_from_doc(doc)

    if not citation_map:
        print("⚠  WARNING: No references found.")
        if not refs_txt:
            print("   Tip: use --refs your_references.txt to supply them externally.")
    else:
        print(f"✅  Found {len(citation_map)} reference(s) via {method}:")
        for k in sorted(citation_map):
            preview = citation_map[k][:75] + ('…' if len(citation_map[k]) > 75 else '')
            print(f"    {k:>3}. {preview}")

    # Process paragraphs (stop at References section)
    replaced = 0
    all_missing = []

    for idx, para in enumerate(doc.paragraphs):
        if ref_start_idx is not None and idx >= ref_start_idx:
            break
        old_xml = para._p.xml
        missing = process_paragraph(para._p, citation_map)
        all_missing.extend(missing)
        if para._p.xml != old_xml:
            replaced += 1

    # Process table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_missing.extend(process_paragraph(para._p, citation_map))

    if all_missing:
        print(f"\n⚠  Citations not found in reference list: {sorted(set(all_missing))}")

    print(f"\n✅  Replaced citations in {replaced} paragraph(s).")
    doc.save(output_path)
    print(f"💾  Saved: {output_path}\n")


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Convert NHSJS citations to Online format.")
    parser.add_argument("input",  help="Input .docx file")
    parser.add_argument("output", help="Output .docx file")
    parser.add_argument("--refs", metavar="FILE.txt",
                        help="Optional plain-text file containing numbered references")
    args = parser.parse_args()
    convert(args.input, args.output, refs_txt=args.refs)