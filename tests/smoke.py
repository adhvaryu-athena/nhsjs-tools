"""
Smoke tests for nhsjs-tools v2.0.0.

Run from the repo root:
    python tests/smoke.py

What this exercises:
  1. LaTeX → Online conversion on tests/fixtures/sample_latex.zip
  2. Standard → Online conversion on tests/fixtures/sample_standard.docx
  3. Revision builder (v2.2 patch model) on the Aaryamann reconstruction
     fixture. The fixture contains:
       - original.docx       (the submitted Standard manuscript)
       - final-tracked.docx  (the human-produced revised tracked-changes
                              version we reconcile against)
       - revision.md         (a v2.2 patch list reverse-engineered from
                              the diff between original and final-tracked)
     Pass conditions: build_all produces 5 outputs, all patches apply
     with FIND-must-match verification, the resulting Standard-Clean.docx
     matches the human final (with changes accepted) at ≥99% character
     level. This is the v2.0.0 regression gate.

Outputs go to a fresh tempdir per run, printed at the start.
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

FIXTURES = Path(__file__).resolve().parent / "fixtures"


def smoke_latex_to_online(out_dir: Path) -> None:
    """Run LaTeX → Online on the sample zip; assert non-empty docx with content."""
    from nhsjs_convert import convert_latex_to_online

    src = FIXTURES / "sample_latex.zip"
    assert src.exists(), f"missing fixture: {src}"

    out_path = out_dir / "latex_to_online.docx"
    convert_latex_to_online(str(src), str(out_path))

    size = out_path.stat().st_size
    assert size > 5000, f"latex output suspiciously small: {size} bytes"

    from docx import Document
    doc = Document(str(out_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "References" in full_text, "latex output missing References heading"
    assert "((" in full_text, "latex output missing any (( online citation"
    print(f"  PASS  latex_to_online    {size:>8} bytes, "
          f"{len(doc.paragraphs)} paras")


def smoke_standard_to_online(out_dir: Path) -> None:
    """Run Standard → Online on the cricket docx; assert citations were converted."""
    from nhsjs_standard_to_online import convert_standard_to_online

    src = FIXTURES / "sample_standard.docx"
    assert src.exists(), f"missing fixture: {src}"

    out_path = out_dir / "standard_to_online.docx"
    convert_standard_to_online(str(src), str(out_path))

    size = out_path.stat().st_size
    assert size > 5000, f"standard output suspiciously small: {size} bytes"

    from docx import Document
    doc = Document(str(out_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "((" in full_text, (
        "standard output has no (( citation brackets — Standard→Online "
        "did not convert any superscript runs"
    )
    print(f"  PASS  standard_to_online {size:>8} bytes, "
          f"{len(doc.paragraphs)} paras")


def smoke_revision_builder(out_dir: Path) -> None:
    """Run the v2.2 patch-application engine against the Aaryamann fixture.

    Asserts the four hard gates from the reconstruction test:
      (a) every patch applies (no FIND mismatches)
      (b) all 5 outputs produced + FILES.md
      (c) OOXML superscript runs present in Standard-Tracked
      (d) Standard-Clean text ≥99% match against human final-tracked
          with changes accepted
    """
    import difflib
    import re

    try:
        import nhsjs_revision_builder as rb
    except ImportError:
        print("  SKIP  revision_builder   (module not built yet)")
        return

    fixture_dir = FIXTURES / "aaryamann"
    revision_md = fixture_dir / "revision.md"
    original_docx = fixture_dir / "original.docx"
    final_tracked = fixture_dir / "final-tracked.docx"

    for p in (revision_md, original_docx, final_tracked):
        assert p.exists(), f"missing Aaryamann fixture: {p}"

    sub = out_dir / "aaryamann_build"
    sub.mkdir(parents=True, exist_ok=True)

    result = rb.build_all(str(revision_md), str(original_docx),
                          None, str(sub))

    # (a) every patch applied
    apply_report = result.get("_apply_report")
    assert apply_report is not None, "build_all did not return _apply_report"
    assert apply_report.failed == 0, (
        f"{apply_report.failed} patches failed to apply (expected 0)"
    )
    n_applied = apply_report.applied

    # (b) all outputs present + non-tiny
    expected_keys = {"standard_tracked", "standard_clean", "online",
                     "response_letter", "student_handoff", "files_md"}
    assert expected_keys <= set(result.keys()), (
        f"missing outputs: {expected_keys - set(result.keys())}"
    )
    for kind in expected_keys:
        path = Path(result[kind])
        assert path.exists() and path.stat().st_size > 100, \
            f"{kind} suspiciously small or missing: {path}"

    # (c) OOXML superscript runs in Standard-Tracked
    from docx import Document
    from docx.oxml.ns import qn
    tracked_doc = Document(result["standard_tracked"])
    sup_count = sum(
        1 for r in tracked_doc.element.body.iter(qn("w:r"))
        for rPr in [r.find(qn("w:rPr"))]
        if rPr is not None
        and rPr.find(qn("w:vertAlign")) is not None
        and rPr.find(qn("w:vertAlign")).get(qn("w:val")) == "superscript"
    )
    assert sup_count >= 30, (
        f"Standard-Tracked has only {sup_count} OOXML superscript runs "
        f"(expected ≥30 for Aaryamann fixture)"
    )

    # (d) Standard-Clean reconciles against human final-tracked (accepted)
    def _accepted_text(p):
        return "".join(
            t.text or "" for t in p._p.iter(qn("w:t"))
            if not any(
                anc.tag == qn("w:del")
                for anc in _iter_ancestors(t)
            )
        )

    def _iter_ancestors(elem):
        a = elem.getparent()
        while a is not None:
            yield a
            a = a.getparent()

    def _norm(s):
        return re.sub(r"\s+", " ", s).strip()

    final_doc = Document(str(final_tracked))
    clean_doc = Document(result["standard_clean"])

    final_paras = [_accepted_text(p) for p in final_doc.paragraphs
                    if _accepted_text(p).strip()]
    clean_paras = [p.text for p in clean_doc.paragraphs if p.text.strip()]

    final_full = "\n".join(_norm(p) for p in final_paras)
    clean_full = "\n".join(_norm(p) for p in clean_paras)
    ratio = difflib.SequenceMatcher(None, clean_full, final_full).ratio()

    assert ratio >= 0.99, (
        f"Standard-Clean reconciliation against human final = {ratio:.3f} "
        f"(expected ≥0.99). v2.0.0 regression — investigate."
    )

    print(f"  PASS  revision_builder    {n_applied} patches applied, "
          f"{sup_count} superscript runs, "
          f"clean↔final ratio {ratio:.3f}")


def main() -> int:
    with tempfile.TemporaryDirectory(prefix="nhsjs_smoke_") as tmp:
        out = Path(tmp)
        print(f"Smoke tests — output dir: {out}")
        smoke_latex_to_online(out)
        smoke_standard_to_online(out)
        smoke_revision_builder(out)
    print("All smoke checks passed.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
